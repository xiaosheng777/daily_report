from __future__ import annotations

from io import BytesIO
from pathlib import Path
from typing import Any

from docx import Document

from src.core.schemas import Report
from src.utils.dates import parse_date

FIELD_NAMES = ("姓名", "部门", "日期", "查重归属日期", "提交时间", "直属负责人", "标题总结", "工作描述")


def build_report_docx(payload: dict[str, Any]) -> bytes:
    document = Document()
    rows = [
        "员工日报", "", "一、基础信息",
        f"姓名：{payload.get('name') or payload.get('employee_name') or ''}",
        f"部门：{payload.get('department') or payload.get('department_name') or ''}",
        f"查重归属日期：{payload.get('report_date') or payload.get('date') or ''}",
        f"提交时间：{payload.get('submitted_at') or payload.get('created_at') or ''}",
        f"直属负责人：{payload.get('manager') or ''}",
        "", "二、今日完成工作",
        f"标题总结：{payload.get('title_summary') or ''}",
        "工作描述：", str(payload.get("work_description") or ""),
    ]
    for line in rows:
        document.add_paragraph(line)
    _append_items(document, "三、测试用例相关工作", payload.get("testcase_items") or [], ["case_id", "用例编号", "描述", "具体步骤", "期望结果"])
    _append_items(document, "四、代码相关工作", payload.get("code_items") or [], ["file_name", "language", "work_type", "description"])
    _append_items(document, "五、文档相关工作", payload.get("document_items") or [], ["document_name", "document_type", "work_type", "description"])
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _append_items(document: Document, title: str, items: list[dict[str, Any]], keys: list[str]) -> None:
    if not any(any(str(v or "").strip() for v in item.values()) for item in items):
        return
    document.add_paragraph("")
    document.add_paragraph(title)
    labels = {
        "scene": "对应场景", "work_type": "工作类型", "owner": "负责人", "case_id": "用例编号",
        "用例编号": "用例编号", "描述": "描述", "具体步骤": "具体步骤", "期望结果": "期望结果",
        "description": "相关说明", "file_name": "文件名称", "language": "语言",
        "document_name": "文档名称", "document_type": "文档类型",
    }
    for idx, item in enumerate(items):
        if idx:
            document.add_paragraph("")
        for key in keys:
            document.add_paragraph(f"{labels.get(key, key)}：{item.get(key) or ''}")


def parse_docx_text(path: str | Path) -> str:
    document = Document(str(path))
    parts: list[str] = []
    parts.extend(p.text.strip() for p in document.paragraphs if p.text.strip())
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def parse_report_docx(file_path: str | Path, owner_user_id: str = "", department_id: str = "") -> Report:
    path = Path(file_path)
    document = Document(str(path))
    lines = [p.text.strip() for p in document.paragraphs if p.text.strip()]
    fields = _extract_fields(lines)
    required = ("姓名", "部门", "直属负责人", "标题总结", "工作描述")
    missing = [name for name in required if not fields.get(name)]
    if not (fields.get("查重归属日期") or fields.get("日期")):
        missing.append("查重归属日期")
    if missing:
        raise ValueError(f"缺少字段: {', '.join(missing)}")
    return Report(
        report_id="", owner_user_id=owner_user_id, employee_name=fields["姓名"], department_id=department_id,
        department_name=fields["部门"], report_date=parse_date(fields.get("查重归属日期") or fields["日期"]), manager=fields["直属负责人"],
        title_summary=fields["标题总结"], work_description=fields["工作描述"], source_type="file", original_filename=path.name,
    )


def _extract_fields(lines: list[str]) -> dict[str, str]:
    fields: dict[str, str] = {}
    current = None
    for line in lines:
        matched = False
        for name in FIELD_NAMES:
            for sep in ("：", ":"):
                prefix = name + sep
                if line.startswith(prefix):
                    current = name
                    fields[name] = line[len(prefix):].strip()
                    matched = True
                    break
            if matched:
                break
        if not matched and current:
            fields[current] = (fields.get(current, "") + " " + line).strip()
    return fields
