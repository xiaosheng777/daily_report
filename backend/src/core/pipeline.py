from __future__ import annotations

import csv
import json
import re
import subprocess
import tempfile
import zipfile
from concurrent.futures import FIRST_COMPLETED, ThreadPoolExecutor, wait
from datetime import datetime, timedelta
from difflib import SequenceMatcher
from io import BytesIO
from pathlib import Path
from typing import Any, Callable

from docx import Document
from openpyxl import Workbook, load_workbook
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity

from src.core.schemas import CheckResult, Finding, Report, ReviewCase, User, duplicate_usage_summary, highest_risk
from src.core.text import low_information, normalize_text, stable_hash
from src.core.testcase_excel import (
    DEFAULT_COMPARE_FIELDS,
    case_key,
    compare_fields_equal,
    differing_fields,
    load_daily_sheet_rows,
)
from src.db.base import DatabaseBackend
from src.llm.base import LLMJudge
from src.storage.file_store import FileStorage
from src.utils.dates import days_between, now_iso, parse_date


class TaskCancelled(RuntimeError):
    """Raised at task checkpoints after a persisted cancellation request."""


class TaskPaused(RuntimeError):
    """Raised after all work already dispatched for the current checkpoint is saved."""


class TaskTerminated(RuntimeError):
    """Raised when the worker requested immediate administrative termination."""


MAX_LEGACY_REPORT_WORKERS = 8
MAX_LOCAL_WORKERS = 32


class DailyReportPipeline:
    def __init__(self, config: dict[str, Any], db: DatabaseBackend, storage: FileStorage, llm_judge: LLMJudge | None = None) -> None:
        self.config = config
        self.db = db
        self.storage = storage
        self.llm_judge = llm_judge
        self.daily_cfg = config.get("daily_duplicate", {})
        self.checks = config.get("checks", {})
        self._code_token_cache: dict[str, tuple[list[str], int]] = {}
        self._artifact_cache: dict[str, list[dict[str, Any]]] = {}
        self._document_text_cache: dict[str, str] = {}
        self._testcase_rows_cache: dict[str, dict[str, dict[str, Any]]] = {}
        self._snapshot_cache: dict[tuple[str, str], dict[str, Any] | None] = {}
        self._task_reports_all: list[Report] = []
        self._normalized_text_by_id: dict[str, str] = {}
        self._text_hash_by_id: dict[str, str] = {}
        self._reports_by_user_date: dict[tuple[str, str], list[Report]] = {}
        self._reports_by_department_date: dict[tuple[str, str], list[Report]] = {}
        self._reports_by_date: dict[str, list[Report]] = {}

    def run(
        self,
        user: User | None = None,
        scope: dict[str, Any] | None = None,
        task_id: str = "",
        result_dir: str = "",
        progress_callback: Callable[[str, int, int, str], None] | None = None,
        case_callback: Callable[[dict[str, Any]], None] | None = None,
        control_callback: Callable[[], str] | None = None,
        resume_cases: list[dict[str, Any]] | None = None,
    ) -> dict[str, Any]:
        scope = scope or {}
        progress = progress_callback or (lambda _stage, _current, _total, _message="": None)
        control = control_callback or (lambda: "")
        self._raise_for_control(control())
        progress("loading", 0, 0, "正在加载查重范围内的日报")
        # Candidate data must stay inside the initiator's organisation scope.
        history_start = self._history_start(str(scope.get("report_date_start") or ""))
        history_end = str(scope.get("report_date_end") or "")
        reports_all = self.db.list_reports(None, start=history_start, end=history_end) if not user or user.role == "admin" else self.db.list_reports(user, start=history_start, end=history_end)
        self._task_reports_all = reports_all
        duplicate_context = self._prepare_duplicate_context(reports_all)
        list_artifacts = getattr(self.db, "list_artifacts_for_reports", None)
        if callable(list_artifacts):
            self._artifact_cache.update(list_artifacts([report.report_id for report in reports_all]))
        current_reports = self._filter_current_reports(reports_all, user, scope)
        completed = {str(case.get("report_id") or ""): case for case in (resume_cases or []) if case.get("report_id")}
        completed = {report_id: case for report_id, case in completed.items() if any(report.report_id == report_id for report in current_reports)}
        progress("preparing", len(completed), len(current_reports), f"已加载 {len(reports_all)} 份候选日报")
        report_cases = self._build_cases_resumable(
            current_reports,
            reports_all,
            completed,
            progress,
            case_callback,
            control,
            duplicate_context,
        )
        self._run_llm_judgement(report_cases, reports_all, progress, case_callback, control)
        missing_reports = self._find_missing_reports(current_reports, user, scope)
        payload = {
            "task_id": task_id,
            "generated_at": now_iso(),
            "summary": {
                "report_count": len(report_cases),
                "high_risk_count": sum(1 for c in report_cases if c.get("overall_risk") == "high"),
                "medium_risk_count": sum(1 for c in report_cases if c.get("overall_risk") == "medium"),
                "unknown_risk_count": sum(1 for c in report_cases if c.get("overall_risk") == "unknown"),
                "triggered_duplicate_checks": sum(1 for c in report_cases if "report_duplicate_check" in (c.get("triggered_checks") or [])),
                "triggered_testcase_checks": sum(1 for c in report_cases if "testcase_check" in (c.get("triggered_checks") or [])),
                "triggered_code_checks": sum(1 for c in report_cases if "code_check" in (c.get("triggered_checks") or [])),
                "triggered_document_checks": sum(1 for c in report_cases if "document_check" in (c.get("triggered_checks") or [])),
                "missing_report_count": len(missing_reports),
                **duplicate_usage_summary(report_cases),
            },
            "report_cases": report_cases,
            "missing_reports": missing_reports,
        }
        if task_id and result_dir:
            self._raise_for_control(control())
            progress("exporting", len(report_cases), len(report_cases), "正在生成任务结果文件")
            self._persist_task_outputs(payload, Path(result_dir), user)
        progress("finished", len(report_cases), len(report_cases), "查重完成")
        return payload

    @staticmethod
    def _raise_for_control(action: str) -> None:
        if action == "pause":
            raise TaskPaused("任务已在安全点暂停")
        if action == "cancel":
            raise TaskCancelled("任务已取消")
        if action == "terminate":
            raise TaskTerminated("任务已强制终止")

    def _build_cases_resumable(
        self,
        current_reports: list[Report],
        all_reports: list[Report],
        completed: dict[str, dict[str, Any]],
        progress: Callable[[str, int, int, str], None],
        case_callback: Callable[[dict[str, Any]], None] | None,
        control: Callable[[], str],
        duplicate_context: tuple[dict[str, int], Any] | None = None,
    ) -> list[dict[str, Any]]:
        remaining = [report for report in current_reports if report.report_id not in completed]
        total = len(current_reports)
        completed_count = len(completed)
        if not remaining:
            return [completed[report.report_id] for report in current_reports if report.report_id in completed]

        self._raise_for_control(control())

        def queue_progress(_stage: str, current: int, _queue_total: int, message: str = "") -> None:
            progress("local_filtering", min(total, completed_count + current), total, message)

        # `_build_cases` checkpoints each completed report before it accepts the
        # next one.  Keep the control callback on the pipeline rather than adding
        # another public/private method argument: a few integrations override the
        # existing helper when testing a task snapshot.
        previous_control = getattr(self, "_local_control", None)
        self._local_control = control
        try:
            cases = self._build_cases(remaining, all_reports, queue_progress, case_callback, duplicate_context)
        finally:
            self._local_control = previous_control
        for case in cases:
            completed[case.report_id] = case.to_dict()
        completed_count += len(cases)
        progress("local_filtering", completed_count, total, f"本地筛选已完成 {completed_count}/{total} 份日报")
        self._raise_for_control(control())
        return [completed[report.report_id] for report in current_reports if report.report_id in completed]

    def _history_start(self, start: str) -> str:
        if not start:
            return ""
        windows = [
            int(self.daily_cfg.get("self_history_days", 30)),
            int(self.daily_cfg.get("cross_user_days", 3)),
            int(self.checks.get("code_history_days", 30)),
            int(self.checks.get("document_history_days", 30)),
            int(self.config.get("testcase_verification", {}).get("self_history_days", 7)),
        ]
        earliest = datetime.strptime(parse_date(start), "%Y-%m-%d").date() - timedelta(days=max(windows))
        return earliest.isoformat()

    def _persist_task_outputs(self, payload: dict[str, Any], result_dir: Path, user: User | None) -> None:
        result_dir.mkdir(parents=True, exist_ok=True)
        (result_dir / "review_cases.json").write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
        self.export_review_xlsx(payload["report_cases"], result_dir / "review_cases.xlsx")
        self.export_review_docx(payload["report_cases"], result_dir / "review_cases.docx")
        self.export_missing_reports_xlsx(payload["missing_reports"], result_dir / "missing_reports.xlsx")
        (result_dir / "run_metadata.json").write_text(json.dumps({"generated_at": payload.get("generated_at"), "config": self.config, "user": user.to_dict() if user else None}, ensure_ascii=False, indent=2), encoding="utf-8")

    def _filter_current_reports(self, reports: list[Report], user: User | None, scope: dict[str, Any]) -> list[Report]:
        start = str(scope.get("report_date_start") or "")
        end = str(scope.get("report_date_end") or "")
        department_scope = str(scope.get("department_scope") or "")
        employee_scope = str(scope.get("employee_scope") or "")
        out = []
        for r in reports:
            if user and user.role == "employee" and r.owner_user_id != user.user_id:
                continue
            if user and user.role == "group_leader" and r.owner_user_id != user.user_id and r.group_id != user.group_id:
                continue
            if user and user.role == "director" and r.department_id != user.department_id:
                continue
            if start and r.report_date < parse_date(start):
                continue
            if end and r.report_date > parse_date(end):
                continue
            if department_scope and department_scope not in {r.department_id, r.department_name}:
                continue
            if employee_scope and employee_scope not in {r.owner_user_id, r.employee_name}:
                continue
            out.append(r)
        return sorted(out, key=lambda r: (r.report_date, r.employee_name))

    def _find_missing_reports(self, reports: list[Report], user: User | None, scope: dict[str, Any]) -> list[dict[str, Any]]:
        start, end = str(scope.get("report_date_start") or ""), str(scope.get("report_date_end") or "")
        if not start or not end:
            return []
        start_date = parse_date(start)
        end_date = parse_date(end)
        if start_date > end_date:
            raise ValueError("开始日期不能晚于结束日期")
        submitters = self.db.list_required_submitters(user)
        department_scope, employee_scope = str(scope.get("department_scope") or ""), str(scope.get("employee_scope") or "")
        submitters = [s for s in submitters if (not department_scope or department_scope in {s.get("department_id"), s.get("department_name")}) and (not employee_scope or employee_scope in {s.get("user_id"), s.get("display_name")})]
        submitted = {(r.owner_user_id, r.report_date) for r in reports}
        days, cursor = [], parse_date(start_date)
        from datetime import datetime
        current = datetime.strptime(cursor, "%Y-%m-%d").date()
        final = datetime.strptime(end_date, "%Y-%m-%d").date()
        while current <= final:
            if current.weekday() < 5:
                days.append(current.isoformat())
            current += timedelta(days=1)
        missing = []
        for submitter in submitters:
            created_date = str(submitter.get("created_at") or "")[:10]
            for report_date in days:
                if created_date and created_date > report_date:
                    continue
                if (submitter["user_id"], report_date) not in submitted:
                    missing.append({"report_date": report_date, "user_id": submitter["user_id"], "employee_name": submitter["display_name"], "role": submitter["role"], "department_id": submitter.get("department_id", ""), "department_name": submitter.get("department_name", ""), "group_id": submitter.get("group_id", ""), "group_name": submitter.get("group_name", "")})
        return missing

    def _build_cases(
        self,
        current_reports: list[Report],
        all_reports: list[Report],
        progress_callback: Callable[[str, int, int, str], None] | None = None,
        case_callback: Callable[[dict[str, Any]], None] | None = None,
        duplicate_context: tuple[dict[str, int], Any] | None = None,
    ) -> list[ReviewCase]:
        progress = progress_callback or (lambda _stage, _current, _total, _message="": None)
        index_by_id, matrix = duplicate_context or self._prepare_duplicate_context(all_reports)
        group_extras_cache: dict[tuple[str, str], list[Finding]] = {}
        cases_by_report_id: dict[str, ReviewCase] = {}
        total = len(current_reports)
        if not current_reports:
            return []
        control = getattr(self, "_local_control", None) or (lambda: "")
        self._raise_for_control(control())
        worker_count = min(total, self._local_worker_count())

        def check(report: Report) -> CheckResult:
            return self._safe_check(
                "report_duplicate_check",
                report,
                lambda: self._daily_duplicate_check(report, all_reports, index_by_id, matrix),
            )

        def build_case(report: Report, duplicate_check: CheckResult) -> ReviewCase:
            checks: list[CheckResult] = [
                self._quality_check(report),
                duplicate_check,
                self._safe_check("testcase_check", report, lambda: self._testcase_check(report, all_reports, group_extras_cache)),
            ]
            if self.checks.get("enable_code_check", True):
                code = self._safe_check("code_check", report, lambda: self._code_token_check(report))
                if code.triggered:
                    checks.append(code)
            if self.checks.get("enable_document_check", True):
                doc = self._safe_check("document_check", report, lambda: self._document_token_check(report))
                if doc.triggered:
                    checks.append(doc)
            triggered = [item.checker_id for item in checks if item.triggered]
            risk = highest_risk(item.risk_level for item in checks if item.triggered) if triggered else "normal"
            return ReviewCase(report.report_id, report.employee_name, report.owner_user_id, report.department_id,
                              report.department_name, report.group_id, report.group_name, report.report_date,
                              report.created_at, report.title_summary, risk, triggered, checks)

        # Keep only `local_worker_count` reports in flight.  This makes a pause or
        # cancellation bounded: completed work is persisted immediately, no large
        # pre-submitted queue keeps running after the control request.
        reports = iter(current_reports)
        futures: dict[Any, Report] = {}
        completed = 0
        stop_action = ""

        with ThreadPoolExecutor(max_workers=worker_count, thread_name_prefix="daily-duplicate-report") as executor:
            def submit_available() -> None:
                while not stop_action and len(futures) < worker_count:
                    try:
                        report = next(reports)
                    except StopIteration:
                        return
                    futures[executor.submit(check, report)] = report

            submit_available()
            while futures:
                done, _ = wait(futures, return_when=FIRST_COMPLETED)
                for future in done:
                    report = futures.pop(future)
                    case = build_case(report, future.result())
                    # This is the Local checkpoint.  It intentionally occurs in
                    # completion order; the returned task payload is reordered
                    # below to preserve the caller's original report order.
                    if case_callback:
                        case_callback(case.to_dict())
                    cases_by_report_id[report.report_id] = case
                    completed += 1
                    progress("checking", completed, total, f"本地筛选已完成 {completed}/{total} 份日报")
                    stop_action = stop_action or control()
                if stop_action == "terminate":
                    # Cancel work that has not started.  The supervisor retains
                    # its existing force-kill behaviour for running threads.
                    for future in futures:
                        future.cancel()
                    raise TaskTerminated("任务已强制终止")
                submit_available()

        self._raise_for_control(stop_action)
        return [cases_by_report_id[report.report_id] for report in current_reports]

    def _prepare_duplicate_context(self, reports: list[Report]) -> tuple[dict[str, int], Any]:
        """Build the immutable duplicate-check data once for the current task."""
        texts = [normalize_text(report.title_summary, report.work_description) for report in reports]
        self._normalized_text_by_id = {report.report_id: text for report, text in zip(reports, texts)}
        self._text_hash_by_id = {report_id: stable_hash(text) for report_id, text in self._normalized_text_by_id.items()}
        self._reports_by_user_date = {}
        self._reports_by_department_date = {}
        self._reports_by_date = {}
        for report in reports:
            self._reports_by_user_date.setdefault((report.owner_user_id, report.report_date), []).append(report)
            self._reports_by_department_date.setdefault((report.department_id, report.report_date), []).append(report)
            self._reports_by_date.setdefault(report.report_date, []).append(report)
        matrix = None
        try:
            if any(texts):
                matrix = TfidfVectorizer(analyzer="char_wb", ngram_range=(2, 5), min_df=1).fit_transform(texts)
        except Exception:
            matrix = None
        return {report.report_id: index for index, report in enumerate(reports)}, matrix

    def _local_worker_count(self) -> int:
        if "local_worker_count" in self.daily_cfg:
            try:
                configured = int(self.daily_cfg["local_worker_count"])
            except (TypeError, ValueError):
                configured = 8
            return max(1, min(MAX_LOCAL_WORKERS, configured))
        try:
            legacy = int(self.daily_cfg.get("report_worker_count", 3))
        except (TypeError, ValueError):
            legacy = 3
        return max(1, min(MAX_LEGACY_REPORT_WORKERS, legacy))

    @staticmethod
    def _safe_check(checker_id: str, report: Report, callback: Callable[[], CheckResult]) -> CheckResult:
        try:
            return callback()
        except TaskCancelled:
            raise
        except Exception as exc:
            summary = f"{checker_id} 执行失败：{exc}"
            return CheckResult(checker_id, True, "failed", "unknown", summary,
                               [Finding(checker_id, "unknown", "execution_failed", report.title_summary, reason=summary)])

    def _quality_check(self, report: Report) -> CheckResult:
        issues = []
        if not report.title_summary.strip():
            issues.append("标题缺失")
        if not report.work_description.strip():
            issues.append("工作描述缺失")
        elif len(report.work_description.replace(" ", "")) < int(self.checks.get("description_too_short_length", 12)):
            issues.append("工作描述过短")
        if low_information(report.title_summary, report.work_description, int(self.checks.get("low_information_min_length", 10))):
            issues.append("信息量偏低")
        if not issues:
            return CheckResult("quality_check", False, "success", "normal", "日报信息量正常")
        risk = "high" if "工作描述缺失" in issues else "medium"
        return CheckResult("quality_check", True, "success", risk, "日报质量需关注：" + "、".join(issues), [Finding("quality_check", risk, "quality", report.title_summary, reason="、".join(issues))])

    def _daily_duplicate_check(self, current: Report, reports: list[Report], index_by_id: dict[str, int], matrix) -> CheckResult:
        candidates = self._collect_candidates(current, reports, index_by_id, matrix)
        threshold = float(self.daily_cfg.get("no_candidate_threshold", 0.55))
        if not candidates or candidates[0]["candidate_score"] < threshold:
            return CheckResult("report_duplicate_check", True, "success", "normal", "未找到足够相似的历史日报")
        top_n = max(1, min(10, int(self.daily_cfg.get("llm_candidate_top_n", 3))))
        llm_threshold = float(self.daily_cfg.get("low_info_candidate_score_threshold", 0.82) if low_information(current.title_summary, current.work_description) else self.daily_cfg.get("llm_candidate_score_threshold", 0.72))
        to_judge = [c for c in candidates[:top_n] if c["candidate_score"] >= llm_threshold or c["hash_same"]]
        if not to_judge:
            to_judge = candidates[:1]
        findings: list[Finding] = []
        for cand in to_judge:
            other: Report = cand["report"]
            risk = "high" if cand["hash_same"] else "medium" if cand["candidate_score"] >= 0.75 else "low"
            reason = "标准化后文本完全一致" if cand["hash_same"] else f"候选相似度 {cand['candidate_score']:.2f}，建议复核"
            details = self._candidate_details(cand)
            requires_llm = cand["candidate_score"] >= llm_threshold or cand["hash_same"]
            if cand["hash_same"]:
                details["llm_called"] = False
                details["decision_source"] = "local_exact_match"
            elif requires_llm and self.llm_judge:
                # Persist this local finding before the LLM phase.  The stable pair
                # id makes the pending queue resumable through task_case_results.
                details["llm_called"] = False
                details["llm_pending"] = True
                details["llm_pair_id"] = self._llm_pair_id(current.report_id, other.report_id)
            else:
                details["llm_called"] = False
                if requires_llm:
                    details["degraded_to_local"] = True
                    details["degraded_reason"] = "llm_unavailable"
                    reason = "大模型未启用或不可用，已使用本地相似度"
            findings.append(Finding("report_duplicate_check", risk, "matched", current.title_summary, f"{other.employee_name}-{other.report_date}-{other.title_summary}", cand["candidate_score"], reason, details))
        risk = highest_risk(f.risk_level for f in findings)
        summary = "；".join(f.reason for f in findings[:3])
        return CheckResult("report_duplicate_check", True, "success", risk, summary, findings)

    def _run_llm_judgement(
        self,
        report_cases: list[dict[str, Any]],
        reports: list[Report],
        progress: Callable[[str, int, int, str], None],
        case_callback: Callable[[dict[str, Any]], None] | None,
        control: Callable[[], str],
    ) -> None:
        """Resolve persisted local duplicate pairs without recomputing local work."""
        reports_by_id = {report.report_id: report for report in reports}
        pending: list[tuple[dict[str, Any], dict[str, Any], dict[str, Any], Report, Report]] = []
        for case in report_cases:
            current = reports_by_id.get(str(case.get("report_id") or ""))
            if not current:
                continue
            duplicate_check = next(
                (check for check in case.get("checks", []) if check.get("checker_id") == "report_duplicate_check"),
                None,
            )
            if not duplicate_check:
                continue
            for finding in duplicate_check.get("findings", []):
                details = finding.get("details") or {}
                if not details.get("llm_pending"):
                    continue
                matched = reports_by_id.get(str(details.get("matched_report_id") or ""))
                if matched:
                    pending.append((case, duplicate_check, finding, current, matched))
                else:
                    details["llm_pending"] = False
                    details["degraded_to_local"] = True
                    details["degraded_reason"] = "matched_report_missing"
                    details["llm_error"] = "matched report is unavailable"
                    finding["reason"] = "候选日报不存在，已使用本地相似度"
                    self._refresh_case_risk(case, duplicate_check)
                    if case_callback:
                        case_callback(case)

        total = len(pending)
        def report_queue_progress() -> None:
            queued = max(0, total - completed - len(futures))
            progress(
                "llm_judging",
                completed,
                total,
                f"LLM 判断：{completed} / {total}（active：{len(futures)}，queued：{queued}）",
            )

        # The progress message is persisted on the task and also provides
        # per-task active/queued telemetry to the resource monitor.
        futures: dict[Any, tuple[dict[str, Any], dict[str, Any], dict[str, Any], Report, Report]] = {}
        completed = 0
        progress("llm_judging", 0, total, f"LLM 判断：0 / {total}（active：0，queued：{total}）")
        if not pending:
            self._raise_for_control(control())
            return
        try:
            configured_workers = int(self.config.get("llm_judge", {}).get("per_task_max_concurrency", 30))
        except (TypeError, ValueError):
            configured_workers = 30
        worker_count = min(total, max(1, configured_workers))
        pending_iter = iter(pending)
        pending_action = ""

        def submit_available(executor: ThreadPoolExecutor) -> None:
            while not pending_action and len(futures) < worker_count:
                try:
                    item = next(pending_iter)
                except StopIteration:
                    return
                case, duplicate_check, finding, current, matched = item
                futures[executor.submit(self._judge_duplicate_pair, current, matched, finding.get("details") or {})] = item

        self._raise_for_control(control())
        with ThreadPoolExecutor(max_workers=worker_count, thread_name_prefix="daily-duplicate-llm") as executor:
            submit_available(executor)
            report_queue_progress()
            while futures:
                done, _ = wait(futures, return_when=FIRST_COMPLETED)
                for future in done:
                    case, duplicate_check, finding, _current, _matched = futures.pop(future)
                    result = future.result()
                    completed += 1
                    pending_action = pending_action or control()
                    self._apply_llm_result(case, duplicate_check, finding, result)
                    if case_callback:
                        case_callback(case)
                submit_available(executor)
                report_queue_progress()
        self._raise_for_control(pending_action or control())

    def _judge_duplicate_pair(self, current: Report, matched: Report, details: dict[str, Any]):
        if not self.llm_judge:
            from src.llm.base import LLMJudgeResult
            return LLMJudgeResult(False, False, error="llm unavailable")
        return self.llm_judge.judge_duplicate(current.to_dict(), matched.to_dict(), self._llm_signals(details))

    def _apply_llm_result(self, case: dict[str, Any], duplicate_check: dict[str, Any], finding: dict[str, Any], result: Any) -> None:
        details = finding.get("details") or {}
        details["llm_pending"] = False
        details["llm_called"] = True
        finding["details"] = details
        if result.ok:
            details["llm_reviewed"] = True
            details["llm_confidence"] = result.confidence
            details["llm_raw"] = result.raw
            finding["risk_level"] = result.duplicate_risk
            finding["reason"] = result.decision_summary or str(finding.get("reason") or "")
        else:
            details["degraded_to_local"] = True
            details["llm_error"] = result.error
            details["degraded_reason"] = "llm_call_failed"
            finding["reason"] = f"大模型调用失败，已使用本地相似度：{result.error}"
        self._refresh_case_risk(case, duplicate_check)

    @staticmethod
    def _llm_pair_id(report_id: str, matched_report_id: str) -> str:
        return f"{report_id}:{matched_report_id}"

    @staticmethod
    def _llm_signals(details: dict[str, Any]) -> dict[str, Any]:
        keys = {
            "match_type", "days_gap", "candidate_score", "text_similarity",
            "semantic_similarity", "recency_score", "hash_same",
        }
        return {key: details[key] for key in keys if key in details}

    @staticmethod
    def _refresh_case_risk(case: dict[str, Any], duplicate_check: dict[str, Any]) -> None:
        findings = list(duplicate_check.get("findings") or [])
        duplicate_check["risk_level"] = highest_risk(finding.get("risk_level", "normal") for finding in findings)
        duplicate_check["summary"] = "；".join(str(finding.get("reason") or "") for finding in findings[:3])
        triggered = [check for check in case.get("checks", []) if check.get("triggered")]
        case["overall_risk"] = highest_risk(check.get("risk_level", "normal") for check in triggered)

    def _collect_candidates(self, current: Report, reports: list[Report], index_by_id: dict[str, int], matrix) -> list[dict[str, Any]]:
        self_days = int(self.daily_cfg.get("self_history_days", 30))
        cross_days = int(self.daily_cfg.get("cross_user_days", 3))
        cross_scope = str(self.daily_cfg.get("cross_user_scope", "department"))
        max_candidates = max(1, int(self.daily_cfg.get("max_candidates", 20)))
        weights = self.daily_cfg.get("score_weights", {})
        wt_text, wt_sem, wt_rec = float(weights.get("text", 0.45)), float(weights.get("semantic", 0.45)), float(weights.get("recency", 0.10))
        current_text = self._normalized_text_by_id.get(current.report_id) or normalize_text(current.title_summary, current.work_description)
        current_hash = self._text_hash_by_id.get(current.report_id) or stable_hash(current_text)
        semantic_scores = None
        if matrix is not None and current.report_id in index_by_id:
            semantic_scores = cosine_similarity(matrix[index_by_id[current.report_id]], matrix).ravel()
        eligible = []
        for other in self._indexed_candidate_reports(current, self_days, cross_days, cross_scope, reports):
            if other.report_id == current.report_id or other.status == "withdrawn":
                continue
            if other.report_date > current.report_date:
                continue
            gap = days_between(current.report_date, other.report_date)
            same_user = other.owner_user_id == current.owner_user_id
            if same_user:
                if gap > self_days:
                    continue
                match_type = "self_repeat"
                horizon = max(1, self_days)
            else:
                if gap > cross_days:
                    continue
                if cross_scope != "all" and other.department_id != current.department_id:
                    continue
                match_type = "cross_user_repeat"
                horizon = max(1, cross_days)
            text = self._normalized_text_by_id.get(other.report_id) or normalize_text(other.title_summary, other.work_description)
            sem = float(semantic_scores[index_by_id[other.report_id]]) if semantic_scores is not None and other.report_id in index_by_id else 0.0
            recency = max(0.0, 1.0 - min(gap, horizon) / horizon)
            hash_same = current_hash == (self._text_hash_by_id.get(other.report_id) or stable_hash(text))
            preliminary = sem * wt_sem + recency * wt_rec + (1.0 if hash_same else 0.0)
            eligible.append((preliminary, other, text, sem, recency, match_type, gap, hash_same))

        # Expensive SequenceMatcher is restricted to the strongest semantic/hash candidates.
        pool_size = max(20, max_candidates)
        eligible.sort(key=lambda item: item[0], reverse=True)
        candidates = []
        for _, other, text, sem, recency, match_type, gap, hash_same in eligible[:pool_size]:
            seq = SequenceMatcher(None, current_text, text).ratio()
            score = round(min(1.0, seq * wt_text + sem * wt_sem + recency * wt_rec), 4)
            candidates.append({"report": other, "candidate_score": score, "text_similarity": round(seq, 4), "semantic_similarity": round(sem, 4), "recency_score": round(recency, 4), "match_type": match_type, "days_gap": gap, "hash_same": hash_same})
        candidates.sort(key=lambda x: (x["hash_same"], x["candidate_score"]), reverse=True)
        return candidates[:max_candidates]

    def _indexed_candidate_reports(
        self,
        current: Report,
        self_days: int,
        cross_days: int,
        cross_scope: str,
        fallback_reports: list[Report],
    ) -> list[Report]:
        if not self._reports_by_date:
            return fallback_reports
        current_day = datetime.strptime(current.report_date, "%Y-%m-%d").date()
        candidates: dict[str, Report] = {}
        for offset in range(max(0, self_days) + 1):
            report_date = (current_day - timedelta(days=offset)).isoformat()
            for report in self._reports_by_user_date.get((current.owner_user_id, report_date), []):
                candidates[report.report_id] = report
        for offset in range(max(0, cross_days) + 1):
            report_date = (current_day - timedelta(days=offset)).isoformat()
            if cross_scope == "all":
                reports_for_day = self._reports_by_date.get(report_date, [])
            else:
                reports_for_day = self._reports_by_department_date.get((current.department_id, report_date), [])
            for report in reports_for_day:
                candidates[report.report_id] = report
        return list(candidates.values())

    @staticmethod
    def _candidate_details(cand: dict[str, Any]) -> dict[str, Any]:
        other: Report = cand["report"]
        return {"matched_report_id": other.report_id, "matched_report_date": other.report_date, "matched_owner_user_id": other.owner_user_id, "match_type": cand["match_type"], "days_gap": cand["days_gap"], "candidate_score": cand["candidate_score"], "text_similarity": cand["text_similarity"], "semantic_similarity": cand["semantic_similarity"], "recency_score": cand["recency_score"], "hash_same": cand["hash_same"]}

    def _testcase_check(
        self,
        report: Report,
        all_reports: list[Report],
        group_extras_cache: dict[tuple[str, str], list[Finding]],
    ) -> CheckResult:
        items = [x for x in report.testcase_items if any(str(v or "").strip() for v in x.values())]
        if not items:
            return CheckResult("testcase_check", False, "skipped", "normal", "未上传测试用例 Excel")
        tc_cfg = self.config.get("testcase_verification", {})
        if not tc_cfg.get("enabled", True):
            return CheckResult("testcase_check", True, "skipped", "unknown", "测试用例核验未启用")
        compare_fields = list(tc_cfg.get("compare_fields") or DEFAULT_COMPARE_FIELDS)
        findings: list[Finding] = []
        findings.extend(self._testcase_self_history_findings(report, items, all_reports, tc_cfg, compare_fields))
        findings.extend(self._testcase_group_consistency_findings(report, items, all_reports, compare_fields, group_extras_cache))
        if not findings:
            return CheckResult("testcase_check", True, "success", "normal", "测试用例核验通过，未发现异常")
        risk = highest_risk(f.risk_level for f in findings)
        summary = "；".join(dict.fromkeys(f.reason for f in findings[:4]))
        return CheckResult("testcase_check", True, "success", risk, summary or "测试用例核验完成", findings)

    def _testcase_self_history_findings(
        self,
        report: Report,
        items: list[dict[str, Any]],
        all_reports: list[Report],
        tc_cfg: dict[str, Any],
        compare_fields: list[str],
    ) -> list[Finding]:
        history_days = int(tc_cfg.get("self_history_days", 7))
        high_rate = float(tc_cfg.get("self_history_high_rate", 0.50))
        medium_rate = float(tc_cfg.get("self_history_medium_rate", 0.20))
        history_rows: dict[str, tuple[dict[str, Any], str]] = {}
        for other in all_reports:
            if other.owner_user_id != report.owner_user_id or other.report_id == report.report_id:
                continue
            if other.status == "withdrawn":
                continue
            if other.report_date >= report.report_date:
                continue
            if days_between(report.report_date, other.report_date) > history_days:
                continue
            for item in other.testcase_items or []:
                cid = case_key(item) or str(item.get("case_id") or "").strip()
                if not cid:
                    continue
                # Prefer the nearest historical day for the same case id.
                prev = history_rows.get(cid)
                if prev is None or other.report_date > prev[1]:
                    history_rows[cid] = (item, other.report_date)

        duplicates: list[Finding] = []
        for item in items:
            cid = case_key(item) or str(item.get("case_id") or "").strip()
            if not cid or cid not in history_rows:
                continue
            hist_item, hist_date = history_rows[cid]
            if not compare_fields_equal(item, hist_item, compare_fields):
                continue
            duplicates.append(
                Finding(
                    "testcase_check",
                    "medium",
                    "self_history_duplicate",
                    cid,
                    f"{hist_date}:{cid}",
                    score=1.0,
                    reason=f"用例 {cid} 与本人 {hist_date} 的用例关键字段一致",
                    details={"case_id": cid, "history_date": hist_date, "match_type": "self_history_duplicate"},
                )
            )

        total = len([x for x in items if case_key(x) or str(x.get("case_id") or "").strip()])
        duplicate_count = len(duplicates)
        rate = (duplicate_count / total) if total else 0.0
        if rate >= high_rate:
            risk = "high"
        elif rate >= medium_rate:
            risk = "medium"
        elif duplicate_count:
            risk = "low"
        else:
            return []
        for finding in duplicates:
            finding.risk_level = risk
            finding.score = round(rate, 4)
            finding.details["duplicate_rate"] = round(rate, 4)
            finding.details["duplicate_count"] = duplicate_count
            finding.details["today_count"] = total
            finding.details["high_rate"] = high_rate
            finding.details["medium_rate"] = medium_rate
        summary_finding = Finding(
            "testcase_check",
            risk,
            "self_history_rate",
            report.employee_name,
            score=round(rate, 4),
            reason=f"本人近 {history_days} 天测试用例查重率 {rate:.0%}（重复 {duplicate_count}/{total}）",
            details={
                "duplicate_rate": round(rate, 4),
                "duplicate_count": duplicate_count,
                "today_count": total,
                "self_history_days": history_days,
                "high_rate": high_rate,
                "medium_rate": medium_rate,
            },
        )
        return [summary_finding, *duplicates]

    def _testcase_group_consistency_findings(
        self,
        report: Report,
        items: list[dict[str, Any]],
        all_reports: list[Report],
        compare_fields: list[str],
        group_extras_cache: dict[tuple[str, str], list[Finding]],
    ) -> list[Finding]:
        if not report.group_id:
            return [
                Finding(
                    "testcase_check",
                    "unknown",
                    "group_missing",
                    report.employee_name,
                    reason="用户未绑定小组，跳过组长汇总一致性检查",
                )
            ]
        snapshot_key = (report.group_id, report.report_date)
        if snapshot_key not in self._snapshot_cache:
            self._snapshot_cache[snapshot_key] = self.db.find_testcase_snapshot(*snapshot_key)
        snapshot = self._snapshot_cache[snapshot_key]
        if not snapshot:
            return [
                Finding(
                    "testcase_check",
                    "medium",
                    "snapshot_missing",
                    report.group_name or report.group_id,
                    reason="该小组当天未上传测试用例汇总表",
                    details={"group_id": report.group_id, "report_date": report.report_date},
                )
            ]
        try:
            leader_rows = self._load_testcase_rows(snapshot["stored_path"])
        except ValueError as exc:
            return [
                Finding(
                    "testcase_check",
                    "high",
                    "sheet_missing",
                    report.group_name or report.group_id,
                    reason=str(exc),
                    details={"snapshot_id": snapshot.get("snapshot_id")},
                )
            ]

        findings: list[Finding] = []
        for item in items:
            cid = case_key(item) or str(item.get("case_id") or "").strip()
            if not cid:
                continue
            leader_row = leader_rows.get(cid)
            if not leader_row:
                findings.append(
                    Finding(
                        "testcase_check",
                        "high",
                        "missing_in_leader",
                        cid,
                        score=0.0,
                        reason=f"组员用例 {cid} 未出现在组长当天汇总表中",
                        details={"case_id": cid, "snapshot_id": snapshot.get("snapshot_id")},
                    )
                )
                continue
            diffs = differing_fields(item, leader_row, compare_fields)
            if diffs:
                findings.append(
                    Finding(
                        "testcase_check",
                        "high",
                        "content_mismatch",
                        cid,
                        score=0.0,
                        reason=f"用例 {cid} 与组长汇总不一致：{('、'.join(diffs))}",
                        details={"case_id": cid, "diff_fields": diffs, "snapshot_id": snapshot.get("snapshot_id")},
                    )
                )

        cache_key = (report.group_id, report.report_date)
        if cache_key not in group_extras_cache:
            member_cases: dict[str, dict[str, Any]] = {}
            for other in all_reports:
                if other.group_id != report.group_id or other.report_date != report.report_date:
                    continue
                if other.status == "withdrawn":
                    continue
                for item in other.testcase_items or []:
                    cid = case_key(item) or str(item.get("case_id") or "").strip()
                    if cid:
                        member_cases[cid] = item
            extras = []
            for cid, leader_row in leader_rows.items():
                if cid.startswith("row_"):
                    continue
                if cid in member_cases:
                    continue
                extras.append(
                    Finding(
                        "testcase_check",
                        "high",
                        "extra_in_leader",
                        cid,
                        score=0.0,
                        reason=f"组长汇总存在组员未提交的用例 {cid}",
                        details={"case_id": cid, "snapshot_id": snapshot.get("snapshot_id"), "group_id": report.group_id},
                    )
                )
            group_extras_cache[cache_key] = extras
        findings.extend(group_extras_cache.get(cache_key, []))
        return findings

    def _load_testcase_rows(self, stored_path: str) -> dict[str, dict[str, Any]]:
        cached = self._testcase_rows_cache.get(stored_path)
        if cached is not None:
            return cached
        path = self.storage.resolve(stored_path)
        rows = load_daily_sheet_rows(path)
        self._testcase_rows_cache[stored_path] = rows
        return rows

    @staticmethod
    def _case_key(row: dict[str, Any]) -> str:
        return case_key(row)

    @staticmethod
    def _row_hash(row: dict[str, Any]) -> str:
        return stable_hash(json.dumps(row, ensure_ascii=False, sort_keys=True))

    def _simple_artifact_check(self, report: Report, artifact_type: str) -> CheckResult:
        checker = "code_check" if artifact_type == "code" else "document_check"
        artifacts = [a for a in self.db.list_artifacts_for_report(report.report_id) if a.get("artifact_type") == artifact_type]
        items = report.code_items if artifact_type == "code" else report.document_items
        if not artifacts and not any(any(str(v or "").strip() for v in item.values()) for item in items):
            return CheckResult(checker, False, "skipped", "normal", "未填写专项条目")
        if artifact_type == "code":
            jplag_path = Path(str(self.config.get("jplag", {}).get("jar_path", "")))
            details = {"jplag_enabled": bool(self.config.get("jplag", {}).get("enabled", True)), "jplag_available": jplag_path.exists(), "artifact_count": len(artifacts), "scope": "self_history_only"}
            summary = "已记录代码附件；JPlag 可用" if jplag_path.exists() else "已记录代码附件；JPlag jar 未提供，正式代码查重会降级"
        else:
            details = {"artifact_count": len(artifacts), "scope": "self_history_only", "docx_tables_supported": True}
            summary = "已记录文档附件，文档查重保持本人历史范围"
        return CheckResult(checker, True, "success", "normal", summary, [Finding(checker, "normal", "recorded", report.title_summary, score=0, reason=summary, details=details)])

    def _code_token_check(self, report: Report) -> CheckResult:
        checker = "code_check"
        artifacts = self._code_artifacts(report)
        items_present = any(any(str(value or "").strip() for value in item.values()) for item in report.code_items)
        if not artifacts and not items_present:
            return CheckResult(checker, False, "skipped", "normal", "未填写专项条目")
        current_tokens, current_files = self._code_tokens(report)
        if not current_tokens:
            summary = "代码 Token 查重未执行：附件中未找到可解析的源代码文件"
            return CheckResult(checker, True, "unknown", "unknown", summary, [Finding(checker, "unknown", "no_source_files", report.title_summary, reason=summary)])
        history_days = int(self.checks.get("code_history_days", 30))
        history = [candidate for candidate in (self._task_reports_all or self.db.list_reports(None)) if candidate.owner_user_id == report.owner_user_id and candidate.report_id != report.report_id and candidate.report_date <= report.report_date and days_between(candidate.report_date, report.report_date) <= history_days]
        history_with_tokens = [(candidate, *self._code_tokens(candidate)) for candidate in history]
        history_with_tokens = [(candidate, tokens, file_count) for candidate, tokens, file_count in history_with_tokens if tokens]
        if not history_with_tokens:
            summary = f"未找到近 {history_days} 天可比对的本人历史代码"
            return CheckResult(checker, True, "success", "normal", summary, [Finding(checker, "normal", "no_history", report.title_summary, reason=summary)])
        matches = []
        for candidate, history_tokens, history_files in history_with_tokens:
            similarity, details = self._code_token_similarity(current_tokens, history_tokens)
            matches.append({"source_name": "当前代码附件：" + "、".join(str(artifact.get("original_filename") or "未命名文件") for artifact in artifacts), "matched_name": f"{candidate.report_date} · {candidate.title_summary}", "similarity": similarity, "history_source_file_count": history_files, **details})
        matches.sort(key=lambda item: item["similarity"], reverse=True)
        similarity = matches[0]["similarity"]
        medium = float(self.checks.get("code_token_similarity_threshold", 0.65))
        high = float(self.checks.get("code_token_high_risk_threshold", 0.85))
        if not 0 <= medium <= high <= 1:
            raise ValueError("代码 Token 查重阈值需满足 0 ≤ 中风险 ≤ 高风险 ≤ 1")
        risk = "high" if similarity >= high else "medium" if similarity >= medium else "normal"
        summary = f"代码 Token 查重完成，最高相似度 {similarity:.1%}（本人近 {history_days} 天历史代码）"
        finding = Finding(checker, risk, "code_token_similarity", report.title_summary, matches[0]["matched_name"], similarity, summary, {"current_source_file_count": current_files, "history_source_file_count": sum(file_count for _, _, file_count in history_with_tokens), "scope": "self_history_only", "engine": "code_token_similarity", "tokenizer": "normalized_code_5gram", "matches": matches})
        return CheckResult(checker, True, "success", risk, summary, [finding])

    def _code_tokens(self, report: Report) -> tuple[list[str], int]:
        cached = self._code_token_cache.get(report.report_id)
        if cached is not None:
            return cached
        allowed = {str(ext).lower() if str(ext).startswith(".") else f".{str(ext).lower()}" for ext in self.config.get("upload", {}).get("allowed_code_ext", [".zip", ".py", ".java", ".js", ".ts", ".cpp", ".c", ".cs", ".go", ".rs", ".php", ".rb", ".html", ".css"])}
        source_extensions = allowed - {".zip"}
        texts: list[str] = []
        for artifact in self._code_artifacts(report):
            name = Path(str(artifact.get("original_filename") or "source")).name
            data = self.storage.read(str(artifact.get("stored_path") or ""))
            if name.lower().endswith(".zip"):
                with zipfile.ZipFile(BytesIO(data)) as archive:
                    members = archive.infolist()
                    self._validate_zip_members(members)
                    for member in members:
                        relative = Path(member.filename)
                        if member.is_dir() or relative.is_absolute() or ".." in relative.parts or relative.suffix.lower() not in source_extensions:
                            continue
                        texts.append(archive.read(member).decode("utf-8", errors="replace"))
            elif Path(name).suffix.lower() in source_extensions:
                texts.append(data.decode("utf-8", errors="replace"))
        tokens = self._normalize_code_tokens("\n".join(texts))
        result = (tokens, len(texts))
        self._code_token_cache[report.report_id] = result
        return result

    @staticmethod
    def _validate_zip_members(members: list[zipfile.ZipInfo]) -> None:
        if len(members) > 2_000:
            raise ValueError("ZIP 文件条目超过 2000 个限制")
        if sum(max(0, int(member.file_size)) for member in members) > 200 * 1024 * 1024:
            raise ValueError("ZIP 解压后数据超过 200 MB 限制")

    @staticmethod
    def _normalize_code_tokens(source: str) -> list[str]:
        without_strings = re.sub(r'"(?:\\.|[^"\\])*"|\'(?:\\.|[^\'\\])*\'|`(?:\\.|[^`\\])*`', " STR ", source)
        without_comments = re.sub(r"/\*[\s\S]*?\*/|//[^\n]*|#[^\n]*", " ", without_strings)
        raw_tokens = re.findall(r"[A-Za-z_]\w*|\d+(?:\.\d+)?|==|!=|<=|>=|=>|\+\+|--|&&|\|\||\S", without_comments)
        keywords = {"and", "as", "async", "await", "break", "case", "catch", "class", "const", "continue", "def", "default", "do", "else", "enum", "except", "export", "extends", "finally", "for", "from", "function", "if", "import", "in", "interface", "let", "new", "not", "null", "package", "private", "protected", "public", "return", "static", "struct", "switch", "this", "throw", "try", "true", "false", "undefined", "var", "void", "while", "with", "yield"}
        normalized = []
        for token in raw_tokens:
            lowered = token.lower()
            if token == "STR": normalized.append("STR")
            elif re.fullmatch(r"\d+(?:\.\d+)?", token): normalized.append("NUM")
            elif re.fullmatch(r"[A-Za-z_]\w*", token): normalized.append(lowered if lowered in keywords else "ID")
            else: normalized.append(token)
        return normalized

    @staticmethod
    def _code_token_similarity(current_tokens: list[str], history_tokens: list[str]) -> tuple[float, dict[str, float | int]]:
        def fingerprints(tokens: list[str]) -> set[tuple[str, ...]]:
            return {tuple(tokens[index:index + 5]) for index in range(max(1, len(tokens) - 4))}
        current_fingerprints, history_fingerprints = fingerprints(current_tokens), fingerprints(history_tokens)
        overlap = len(current_fingerprints & history_fingerprints) / len(current_fingerprints | history_fingerprints) if current_fingerprints or history_fingerprints else 0.0
        sequence = SequenceMatcher(None, current_tokens, history_tokens, autojunk=False).ratio()
        return 0.7 * overlap + 0.3 * sequence, {"current_token_count": len(current_tokens), "matched_token_count": len(current_fingerprints & history_fingerprints), "token_overlap": overlap, "sequence_similarity": sequence}

    def _batch_code_jplag_checks(self, current_reports: list[Report], all_reports: list[Report]) -> dict[str, CheckResult]:
        """Run JPlag once per owner when a task contains several code submissions."""
        jplag = self.config.get("jplag", {})
        jar_path = Path(str(jplag.get("jar_path", "")))
        if not jplag.get("enabled", True) or not jar_path.is_file():
            return {}
        reports_with_code = [report for report in current_reports if self._code_artifacts(report)]
        grouped: dict[str, list[Report]] = {}
        for report in reports_with_code:
            grouped.setdefault(report.owner_user_id, []).append(report)
        results: dict[str, CheckResult] = {}
        history_days = int(self.checks.get("code_history_days", 30))
        for owner_id, targets in grouped.items():
            if len(targets) < 2:
                continue
            earliest = min(report.report_date for report in targets)
            latest = max(report.report_date for report in targets)
            target_ids = {report.report_id for report in targets}
            candidates = [
                report for report in all_reports
                if report.owner_user_id == owner_id and self._code_artifacts(report)
                and (report.report_id in target_ids or (report.report_date <= latest and days_between(report.report_date, earliest) <= history_days))
            ]
            try:
                results.update(self._run_batched_code_jplag(jar_path, targets, candidates, history_days))
            except (OSError, subprocess.SubprocessError, ValueError, zipfile.BadZipFile) as exc:
                for report in targets:
                    summary = f"代码 JPlag 批量查重执行失败：{exc}"
                    results[report.report_id] = CheckResult("code_check", True, "unknown", "unknown", summary, [Finding("code_check", "unknown", "execution_failed", report.title_summary, reason=summary)])
        return results

    def _run_batched_code_jplag(self, jar_path: Path, targets: list[Report], candidates: list[Report], history_days: int) -> dict[str, CheckResult]:
        jplag = self.config.get("jplag", {})
        medium = float(jplag.get("code_similarity_threshold", 0.65))
        high = float(jplag.get("code_high_risk_threshold", 0.85))
        if not 0 <= medium <= high <= 1:
            raise ValueError("JPlag 代码阈值需满足 0 ≤ 中风险 ≤ 高风险 ≤ 1")
        with tempfile.TemporaryDirectory(prefix="daily-report-jplag-") as work:
            root = Path(work)
            new_root, old_root = root / "new", root / "old"
            report_by_submission: dict[str, Report] = {}
            source_files: dict[str, int] = {}
            target_ids = {report.report_id for report in targets}
            for index, report in enumerate(candidates):
                submission_name = f"submission_{index:03d}"
                file_count = self._write_code_submission(old_root / submission_name, self._code_artifacts(report))
                if file_count:
                    report_by_submission[submission_name] = report
                    source_files[report.report_id] = file_count
                    if report.report_id in target_ids:
                        self._write_code_submission(new_root / submission_name, self._code_artifacts(report))
            if not target_ids.issubset(source_files):
                return {}
            _, raw_matches = self._run_jplag_result(jar_path, root, new_root, old_root, str(jplag.get("code_language", "multi")), int(jplag.get("code_min_tokens", 12)), int(jplag.get("code_timeout_seconds", 120)), ["multi", "--languages", str(jplag.get("code_multi_languages", "c,cpp,csharp,go,java,javascript,kotlin,python3,rust,typescript"))] if str(jplag.get("code_language", "multi")) == "multi" else [])

        match_rows: dict[str, list[dict[str, Any]]] = {report_id: [] for report_id in target_ids}
        for match in raw_matches:
            left_name, right_name = str(match.get("source") or ""), str(match.get("matched") or "")
            left_report = next((report for name, report in report_by_submission.items() if name in left_name), None)
            right_report = next((report for name, report in report_by_submission.items() if name in right_name), None)
            if not left_report or not right_report or left_report.report_id == right_report.report_id:
                continue
            for current, historical in ((left_report, right_report), (right_report, left_report)):
                if current.report_id not in target_ids or historical.report_date > current.report_date or days_between(historical.report_date, current.report_date) > history_days:
                    continue
                match_rows[current.report_id].append({
                    "source_name": "当前代码附件：" + "、".join(str(artifact.get("original_filename") or "未命名文件") for artifact in self._code_artifacts(current)),
                    "matched_name": f"{historical.report_date} · {historical.title_summary}",
                    "similarity": float(match.get("similarity") or 0),
                })

        results = {}
        for report in targets:
            valid_history = [candidate for candidate in candidates if candidate.report_id != report.report_id and candidate.report_date <= report.report_date and days_between(candidate.report_date, report.report_date) <= history_days]
            if not valid_history:
                summary = f"未找到近 {history_days} 天可比对的本人历史代码"
                results[report.report_id] = CheckResult("code_check", True, "success", "normal", summary, [Finding("code_check", "normal", "no_history", report.title_summary, reason=summary)])
                continue
            matches = sorted(match_rows[report.report_id], key=lambda item: item["similarity"], reverse=True)
            similarity = matches[0]["similarity"] if matches else 0.0
            risk = "high" if similarity >= high else "medium" if similarity >= medium else "normal"
            history_files = sum(source_files.get(candidate.report_id, 0) for candidate in valid_history)
            summary = f"JPlag 代码批量查重完成，最高相似度 {similarity:.1%}（本人近 {history_days} 天历史代码）"
            finding = Finding("code_check", risk, "jplag_code", report.title_summary, matches[0]["matched_name"] if matches else "本人历史代码", similarity, summary, {
                "language": str(jplag.get("code_language", "multi")), "current_source_file_count": source_files[report.report_id],
                "history_source_file_count": history_files, "scope": "self_history_only", "engine": "jplag", "execution": "batched",
                "matches": matches,
            })
            results[report.report_id] = CheckResult("code_check", True, "success", risk, summary, [finding])
        return results

    def _code_jplag_check(self, report: Report) -> CheckResult:
        checker = "code_check"
        artifacts = self._code_artifacts(report)
        items_present = any(any(str(value or "").strip() for value in item.values()) for item in report.code_items)
        if not artifacts and not items_present:
            return CheckResult(checker, False, "skipped", "normal", "未填写专项条目")
        if not artifacts:
            return CheckResult(checker, True, "success", "normal", "已记录代码条目；未上传可供 JPlag 查重的代码附件")

        jplag = self.config.get("jplag", {})
        jar_path = Path(str(jplag.get("jar_path", "")))
        if not jplag.get("enabled", True) or not jar_path.is_file():
            summary = "代码 JPlag 查重未执行：JPlag jar 不可用"
            return CheckResult(checker, True, "unknown", "unknown", summary, [Finding(checker, "unknown", "unavailable", report.title_summary, reason=summary)])

        history_days = int(self.checks.get("code_history_days", 30))
        history = [candidate for candidate in (self._task_reports_all or self.db.list_reports(None)) if candidate.owner_user_id == report.owner_user_id and candidate.report_id != report.report_id and candidate.report_date <= report.report_date and days_between(candidate.report_date, report.report_date) <= history_days]
        history_with_code = [(candidate, self._code_artifacts(candidate)) for candidate in history]
        history_with_code = [(candidate, submission) for candidate, submission in history_with_code if submission]
        if not history_with_code:
            summary = f"未找到近 {history_days} 天可比对的本人历史代码"
            return CheckResult(checker, True, "success", "normal", summary, [Finding(checker, "normal", "no_history", report.title_summary, reason=summary)])

        try:
            jplag_result = self._run_code_jplag(jar_path, artifacts, [submission for _, submission in history_with_code], True)
            similarity, current_files, history_files = jplag_result[:3]
            raw_matches = jplag_result[3] if len(jplag_result) > 3 else []
        except (OSError, subprocess.SubprocessError, ValueError, zipfile.BadZipFile) as exc:
            summary = f"代码 JPlag 查重执行失败：{exc}"
            return CheckResult(checker, True, "unknown", "unknown", summary, [Finding(checker, "unknown", "execution_failed", report.title_summary, reason=summary)])
        if not current_files or not history_files:
            summary = "代码 JPlag 查重未执行：附件中未找到受支持的源代码文件"
            return CheckResult(checker, True, "unknown", "unknown", summary, [Finding(checker, "unknown", "no_source_files", report.title_summary, reason=summary)])

        medium = float(jplag.get("code_similarity_threshold", 0.65))
        high = float(jplag.get("code_high_risk_threshold", 0.85))
        if not 0 <= medium <= high <= 1:
            raise ValueError("JPlag 代码阈值需满足 0 ≤ 中风险 ≤ 高风险 ≤ 1")
        risk = "high" if similarity >= high else "medium" if similarity >= medium else "normal"
        summary = f"JPlag 代码查重完成，最高相似度 {similarity:.1%}（本人近 {history_days} 天历史代码）"
        finding = Finding(checker, risk, "jplag_code", report.title_summary, "本人历史代码", similarity, summary, {
            "language": str(jplag.get("code_language", "multi")), "current_source_file_count": current_files,
            "history_source_file_count": history_files, "scope": "self_history_only", "engine": "jplag",
            "matches": self._format_jplag_matches(
                raw_matches,
                "当前代码附件：" + "、".join(str(a.get("original_filename") or "未命名文件") for a in artifacts),
                [f"{candidate.report_date} · {candidate.title_summary}" for candidate, _ in history_with_code],
            ),
        })
        return CheckResult(checker, True, "success", risk, summary, [finding])

    def _code_artifacts(self, report: Report) -> list[dict[str, Any]]:
        return [artifact for artifact in self._artifacts(report.report_id) if artifact.get("artifact_type") == "code"]

    def _document_token_check(self, report: Report) -> CheckResult:
        checker = "document_check"
        artifacts = self._document_artifacts(report)
        items_present = any(any(str(value or "").strip() for value in item.values()) for item in report.document_items)
        if not artifacts and not items_present:
            return CheckResult(checker, False, "skipped", "normal", "未填写专项条目")
        if not artifacts:
            return CheckResult(checker, True, "success", "normal", "已记录文档条目；未上传可供正文 Token 查重的文档附件")

        history_days = int(self.checks.get("document_history_days", 30))
        history = [candidate for candidate in (self._task_reports_all or self.db.list_reports(None)) if candidate.owner_user_id == report.owner_user_id and candidate.report_id != report.report_id and candidate.report_date <= report.report_date and days_between(candidate.report_date, report.report_date) <= history_days]
        historical_artifacts = [(candidate, artifact) for candidate in history for artifact in self._document_artifacts(candidate)]
        if not historical_artifacts:
            summary = f"未找到近 {history_days} 天可比对的本人历史文档"
            return CheckResult(checker, True, "success", "normal", summary, [Finding(checker, "normal", "no_history", report.title_summary, reason=summary)])

        try:
            current_documents = [(artifact, self._extract_document_text(artifact)) for artifact in artifacts]
            history_documents = [(history_report, artifact, self._extract_document_text(artifact)) for history_report, artifact in historical_artifacts]
        except ValueError as exc:
            summary = f"文档 Token 查重未执行：{exc}"
            return CheckResult(checker, True, "unknown", "unknown", summary, [Finding(checker, "unknown", "extract_failed", report.title_summary, reason=summary)])
        current_documents = [(artifact, text) for artifact, text in current_documents if text.strip()]
        history_documents = [(history_report, artifact, text) for history_report, artifact, text in history_documents if text.strip()]
        if not current_documents or not history_documents:
            summary = "文档 Token 查重未执行：没有可提取的正文（支持 docx、pdf、txt、md、xlsx）"
            return CheckResult(checker, True, "unknown", "unknown", summary, [Finding(checker, "unknown", "no_extractable_text", report.title_summary, reason=summary)])

        matches = []
        for current_artifact, current_text in current_documents:
            for history_report, history_artifact, history_text in history_documents:
                similarity, token_details = self._document_token_similarity(current_text, history_text)
                matches.append({
                    "source_name": str(current_artifact.get("original_filename") or "未命名文件"),
                    "matched_name": f"{history_report.report_date} · {history_report.title_summary} · {history_artifact.get('original_filename') or '未命名文件'}",
                    "similarity": similarity,
                    **token_details,
                })
        matches.sort(key=lambda item: item["similarity"], reverse=True)
        similarity = matches[0]["similarity"]
        medium = float(self.checks.get("document_token_similarity_threshold", 0.65))
        high = float(self.checks.get("document_token_high_risk_threshold", 0.85))
        if not 0 <= medium <= high <= 1:
            raise ValueError("文档 Token 查重阈值需满足 0 ≤ 中风险 ≤ 高风险 ≤ 1")
        risk = "high" if similarity >= high else "medium" if similarity >= medium else "normal"
        summary = f"文档 Token 查重完成，最高相似度 {similarity:.1%}（本人近 {history_days} 天历史文档）"
        finding = Finding(checker, risk, "token_similarity", report.title_summary, matches[0]["matched_name"], similarity, summary, {
            "current_document_count": len(current_documents), "history_document_count": len(history_documents),
            "scope": "self_history_only", "engine": "token_similarity", "tokenizer": "han_character_and_word",
            "matches": matches,
        })
        return CheckResult(checker, True, "success", risk, summary, [finding])

    def _document_artifacts(self, report: Report) -> list[dict[str, Any]]:
        return [artifact for artifact in self._artifacts(report.report_id) if artifact.get("artifact_type") == "document"]

    def _artifacts(self, report_id: str) -> list[dict[str, Any]]:
        if report_id not in self._artifact_cache:
            self._artifact_cache[report_id] = self.db.list_artifacts_for_report(report_id)
        return self._artifact_cache[report_id]

    def _extract_document_texts(self, artifacts: list[dict[str, Any]]) -> list[str]:
        texts = []
        for artifact in artifacts:
            text = self._extract_document_text(artifact)
            if text.strip():
                texts.append(text)
        return texts

    def _extract_document_text(self, artifact: dict[str, Any]) -> str:
        cache_key = str(artifact.get("artifact_id") or artifact.get("stored_path") or "")
        if cache_key in self._document_text_cache:
            return self._document_text_cache[cache_key]
        name = str(artifact.get("original_filename") or "").lower()
        data = self.storage.read(str(artifact.get("stored_path") or ""))
        max_bytes = int(self.config.get("upload", {}).get("max_file_mb", 50)) * 1024 * 1024
        if len(data) > max_bytes:
            raise ValueError(f"文档超过 {max_bytes // 1024 // 1024} MB 限制")
        if name.endswith(".docx"):
            document = Document(BytesIO(data))
            parts = [paragraph.text for paragraph in document.paragraphs]
            parts.extend(cell.text for table in document.tables for row in table.rows for cell in row.cells)
            text = "\n".join(parts)
        elif name.endswith(".pdf"):
            try:
                from pypdf import PdfReader
            except ImportError as exc:
                raise ValueError("缺少 pypdf 依赖，无法读取 PDF") from exc
            pages = PdfReader(BytesIO(data)).pages
            if len(pages) > 500:
                raise ValueError("PDF 超过 500 页限制")
            text = "\n".join(page.extract_text() or "" for page in pages)
        elif name.endswith((".txt", ".md")):
            text = data.decode("utf-8", errors="replace")
        elif name.endswith(".xlsx"):
            workbook = load_workbook(BytesIO(data), read_only=True, data_only=True)
            text = "\n".join(str(cell) for sheet in workbook.worksheets for row in sheet.iter_rows(values_only=True) for cell in row if cell not in (None, ""))
        else:
            text = ""
        if len(text) > 2_000_000:
            raise ValueError("文档正文超过 200 万字符限制")
        self._document_text_cache[cache_key] = text
        return text

    @staticmethod
    def _document_token_similarity(current_text: str, history_text: str) -> tuple[float, dict[str, float | int]]:
        current_tokens = re.findall(r"[\u3400-\u4dbf\u4e00-\u9fff\uf900-\ufaff]|[a-zA-Z0-9_]+", normalize_text(current_text))
        history_tokens = re.findall(r"[\u3400-\u4dbf\u4e00-\u9fff\uf900-\ufaff]|[a-zA-Z0-9_]+", normalize_text(history_text))
        if not current_tokens or not history_tokens:
            return 0.0, {"current_token_count": len(current_tokens), "matched_token_count": 0, "token_overlap": 0.0, "sequence_similarity": 0.0}
        current_set, history_set = set(current_tokens), set(history_tokens)
        overlap = len(current_set & history_set) / len(current_set | history_set)
        sequence = SequenceMatcher(None, current_tokens, history_tokens, autojunk=False).ratio()
        score = 0.55 * overlap + 0.45 * sequence
        return score, {
            "current_token_count": len(current_tokens),
            "matched_token_count": len(current_set & history_set),
            "token_overlap": overlap,
            "sequence_similarity": sequence,
        }

    def _run_code_jplag(self, jar_path: Path, current_artifacts: list[dict[str, Any]], history_submissions: list[list[dict[str, Any]]], include_matches: bool = False) -> tuple[float, int, int] | tuple[float, int, int, list[dict[str, Any]]]:
        jplag = self.config.get("jplag", {})
        with tempfile.TemporaryDirectory(prefix="daily-report-jplag-") as work:
            root = Path(work)
            new_root, old_root = root / "new", root / "old"
            current_files = self._write_code_submission(new_root / "current", current_artifacts)
            history_files = sum(self._write_code_submission(old_root / f"history_{index:03d}", artifacts) for index, artifacts in enumerate(history_submissions))
            if not current_files or not history_files:
                return 0.0, current_files, history_files
            language = str(jplag.get("code_language", "multi"))
            multi_languages = str(jplag.get("code_multi_languages", "c,cpp,csharp,go,java,javascript,kotlin,python3,rust,typescript"))
            language_options = ["multi", "--languages", multi_languages] if language == "multi" else []
            similarity, matches = self._run_jplag_result(jar_path, root, new_root, old_root, language, int(jplag.get("code_min_tokens", 12)), int(jplag.get("code_timeout_seconds", 120)), language_options)
            return (similarity, current_files, history_files, matches) if include_matches else (similarity, current_files, history_files)

    def _run_jplag(self, jar_path: Path, root: Path, new_root: Path, old_root: Path, language: str, min_tokens: int, timeout: int, language_options: list[str] | None = None) -> float:
        return self._run_jplag_result(jar_path, root, new_root, old_root, language, min_tokens, timeout, language_options)[0]

    def _run_jplag_result(self, jar_path: Path, root: Path, new_root: Path, old_root: Path, language: str, min_tokens: int, timeout: int, language_options: list[str] | None = None) -> tuple[float, list[dict[str, Any]]]:
        jplag = self.config.get("jplag", {})
        result_file = root / "result"
        # JPlag's default AUTO mode may select RUN_AND_VIEW and open its report viewer.
        # RUN keeps it headless: it only writes result files for this application to parse.
        command = [str(jplag.get("java_command", "java")), "-jar", str(jar_path), "--mode", "RUN", "--language", language, "--min-tokens", str(min_tokens), "--csv-export", "--overwrite", "--result-file", str(result_file), "--new", str(new_root), "--old", str(old_root), *(language_options or [])]
        completed = subprocess.run(command, cwd=root, capture_output=True, text=True, timeout=timeout, check=False)
        if completed.returncode:
            message = (completed.stderr or completed.stdout or "JPlag returned a non-zero exit code").strip()
            raise ValueError(message if len(message) <= 1000 else f"{message[:500]}\n…\n{message[-500:]}")
        similarity, matches = self._read_jplag_result(root)
        if similarity is None:
            raise ValueError("JPlag 未生成可解析的相似度结果")
        return similarity, matches

    def _write_code_submission(self, submission: Path, artifacts: list[dict[str, Any]]) -> int:
        allowed = {str(ext).lower() if str(ext).startswith(".") else f".{str(ext).lower()}" for ext in self.config.get("upload", {}).get("allowed_code_ext", [".zip", ".py", ".java", ".js", ".ts", ".cpp", ".c", ".cs", ".go", ".rs", ".php", ".rb", ".html", ".css"])}
        jplag_source_extensions = {".c", ".cpp", ".cc", ".cxx", ".cs", ".go", ".java", ".js", ".kt", ".kts", ".py", ".rs", ".scala", ".scm", ".ss", ".ts", ".tsx"}
        source_extensions = allowed.intersection(jplag_source_extensions)
        count = 0
        for index, artifact in enumerate(artifacts):
            name = Path(str(artifact.get("original_filename") or "source")).name
            data = self.storage.read(str(artifact.get("stored_path") or ""))
            if name.lower().endswith(".zip"):
                with zipfile.ZipFile(BytesIO(data)) as archive:
                    members = archive.infolist()
                    self._validate_zip_members(members)
                    for member in members:
                        relative = Path(member.filename)
                        if member.is_dir() or relative.is_absolute() or ".." in relative.parts or relative.suffix.lower() not in source_extensions:
                            continue
                        target = submission / f"archive_{index:03d}" / relative
                        target.parent.mkdir(parents=True, exist_ok=True)
                        target.write_bytes(archive.read(member))
                        count += 1
            elif Path(name).suffix.lower() in source_extensions:
                submission.mkdir(parents=True, exist_ok=True)
                (submission / f"file_{index:03d}_{name}").write_bytes(data)
                count += 1
        return count

    @staticmethod
    def _read_jplag_similarity(root: Path) -> float | None:
        return DailyReportPipeline._read_jplag_result(root)[0]

    @staticmethod
    def _read_jplag_result(root: Path) -> tuple[float | None, list[dict[str, Any]]]:
        scores = []
        matches = []
        for path in root.rglob("*.csv"):
            with path.open("r", encoding="utf-8-sig", newline="") as handle:
                for row in csv.DictReader(handle):
                    similarity_values = []
                    for key, value in row.items():
                        if "similarity" not in str(key).lower() or value in (None, ""):
                            continue
                        try:
                            score = float(str(value).strip().rstrip("%"))
                        except ValueError:
                            continue
                        normalized = score / 100 if str(value).strip().endswith("%") or score > 1 else score
                        scores.append(normalized)
                        similarity_values.append(normalized)
                    submission_names = [str(value).strip() for key, value in row.items() if value not in (None, "") and "submission" in str(key).lower()]
                    normalized_names = [name.replace("\\", "/").lower() for name in submission_names]
                    if len(submission_names) >= 2 and similarity_values and any("new/" in name or "old/" in name for name in normalized_names[:2]):
                        matches.append({"source": submission_names[0], "matched": submission_names[1], "similarity": max(similarity_values)})
        if scores:
            return max(0.0, min(1.0, max(scores))), matches
        for path in root.rglob("*.zip"):
            with zipfile.ZipFile(path) as archive:
                for name in archive.namelist():
                    if not name.endswith(".json"):
                        continue
                    data = json.loads(archive.read(name).decode("utf-8"))
                    DailyReportPipeline._collect_similarity_values(data, scores)
        return (max(0.0, min(1.0, max(scores))), matches) if scores else (None, [])

    @staticmethod
    def _format_jplag_matches(matches: list[dict[str, Any]], current_label: str, history_labels: list[str]) -> list[dict[str, Any]]:
        formatted = []
        for match in matches:
            source, matched = str(match.get("source") or ""), str(match.get("matched") or "")
            source_is_current = "/new/" in source.replace("\\", "/") or "new_" in source
            matched_is_current = "/new/" in matched.replace("\\", "/") or "new_" in matched
            if matched_is_current and not source_is_current:
                source, matched = matched, source
            history_index = next((index for index in range(len(history_labels)) if f"history_{index:03d}" in matched or f"old_{index:03d}" in matched), None)
            formatted.append({
                "source_name": current_label if source_is_current or matched_is_current else source,
                "matched_name": history_labels[history_index] if history_index is not None else matched,
                "similarity": max(0.0, min(1.0, float(match.get("similarity") or 0))),
            })
        return sorted(formatted, key=lambda item: item["similarity"], reverse=True)

    @staticmethod
    def _collect_similarity_values(value: Any, scores: list[float], key: str = "") -> None:
        if isinstance(value, dict):
            for child_key, child in value.items():
                DailyReportPipeline._collect_similarity_values(child, scores, str(child_key))
        elif isinstance(value, list):
            for child in value:
                DailyReportPipeline._collect_similarity_values(child, scores, key)
        elif "similarity" in key.lower() and isinstance(value, (int, float)):
            scores.append(float(value) / 100 if float(value) > 1 else float(value))

    @staticmethod
    def build_report_id(report: Report) -> str:
        return "R_" + stable_hash(f"{report.owner_user_id}|{report.report_date}|{report.title_summary}|{report.work_description}")[:16]

    @staticmethod
    def export_review_xlsx(cases: list[dict[str, Any]], path: Path) -> None:
        wb = Workbook()
        ws = wb.active
        ws.title = "review_cases"
        ws.append(["查重归属日期", "提交时间", "员工", "部门", "标题", "总体风险", "触发检查", "摘要"])
        for c in cases:
            summary = "；".join(check.get("summary", "") for check in c.get("checks", []) if check.get("triggered"))
            ws.append([c.get("report_date"), c.get("submitted_at"), c.get("employee_name"), c.get("department_name"), c.get("title_summary"), c.get("overall_risk"), ",".join(c.get("triggered_checks", [])), summary])
        path.parent.mkdir(parents=True, exist_ok=True)
        wb.save(path)

    @staticmethod
    def export_review_docx(cases: list[dict[str, Any]], path: Path) -> None:
        doc = Document()
        doc.add_heading("查重任务报告", level=1)
        doc.add_paragraph(f"共 {len(cases)} 份日报")
        for c in cases:
            doc.add_heading(f"{c.get('report_date')} · {c.get('employee_name')} · {c.get('overall_risk')}", level=2)
            doc.add_paragraph(f"查重归属日期：{c.get('report_date') or ''}；提交时间：{c.get('submitted_at') or ''}")
            doc.add_paragraph(str(c.get("title_summary") or ""))
            for check in c.get("checks", []):
                if check.get("triggered"):
                    doc.add_paragraph(f"{check.get('checker_id')}: {check.get('summary')}")
                    for f in check.get("findings", []):
                        doc.add_paragraph(f"- {f.get('risk_level')} | {f.get('reason')}")
        path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(path)

    @staticmethod
    def export_missing_reports_xlsx(rows: list[dict[str, Any]], path: Path) -> None:
        wb = Workbook()
        ws = wb.active
        ws.title = "missing_reports"
        ws.append(["日期", "姓名", "职务", "部门", "小组"])
        for row in rows:
            ws.append([row.get("report_date"), row.get("employee_name"), row.get("role"), row.get("department_name"), row.get("group_name")])
        path.parent.mkdir(parents=True, exist_ok=True)
        wb.save(path)
