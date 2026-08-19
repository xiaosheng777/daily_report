from __future__ import annotations

import argparse
import csv
import io
import json
import mimetypes
import shutil
import tempfile
import time
from datetime import datetime, timedelta
from email.parser import BytesParser
from email.policy import default
from http import HTTPStatus
from http.server import BaseHTTPRequestHandler, ThreadingHTTPServer
from pathlib import Path
from urllib.parse import parse_qs, quote, unquote, urlparse

from src.core.docx_io import build_report_docx
from src.core.pipeline import DailyReportPipeline
from src.core.schemas import Report, duplicate_usage_summary
from src.db.factory import create_database
from src.db.sqlite_backend import (
    BUSINESS_ROLES,
    ROLE_DIRECTOR,
    ROLE_GROUP_LEADER,
    ROLE_MINISTER,
    can_admin,
    can_manage,
)
from src.llm.factory import create_llm_judge
from src.monitoring.diagnostics import build_history as build_resource_history
from src.monitoring.diagnostics import build_snapshot as build_resource_snapshot
from src.storage.file_store import FileStorage
from src.utils.config import default_config_path, load_config
from src.utils.dates import SHANGHAI_TIMEZONE, now_iso, parse_date, ym
from src.utils.security import token
from src.utils.runtime_log import runtime_log, runtime_log_json

SESSION_COOKIE = "daily_report_session"
REQUIRED_REPORT_FIELDS = ("title_summary", "work_description")
RESET_CONFIRMATION_PHRASE = "清除全部业务数据"
AUDIT_CATEGORIES = {"authentication", "administration", "business", "export"}
AUDIT_SEVERITIES = {"info", "warning", "critical"}
AUDIT_OUTCOMES = {"success", "failure"}
AUDIT_SENSITIVE_PARTS = ("password", "api_key", "secret", "token", "cookie", "authorization")
RUNTIME_LOG_LEVELS = {"info", "warning", "error"}
RUNTIME_LOG_SOURCES = {"web", "worker", "scheduler", "task", "llm", "monitor"}


class RequestTooLarge(ValueError):
    pass


def attachment_header(download_name: str) -> str:
    safe_name = Path(download_name).name.replace('"', "")
    ascii_name = safe_name.encode("ascii", "ignore").decode("ascii") or "download"
    return f'attachment; filename="{ascii_name}"; filename*=UTF-8\'\'{quote(safe_name, safe="")}'


def flatten_settings(config: dict) -> dict:
    dd = config.get("daily_duplicate", {})
    submission = config.get("daily_report_submission", {})
    automatic = config.get("automatic_daily_duplicate", {})
    weights = dd.get("score_weights", {})
    checks = config.get("checks", {})
    llm = config.get("llm_judge", {})
    task_runner = config.get("task_runner", {})
    tc = config.get("testcase_verification", {})
    return {
        "daily_report_submission.rollover_time": str(submission.get("rollover_time", submission.get("cutoff_time", "09:00"))),
        "automatic_daily_duplicate.enabled": bool(automatic.get("enabled", True)),
        "automatic_daily_duplicate.run_at": str(automatic.get("run_at", "12:00")),
        "task_runner.max_concurrent_tasks": int(task_runner.get("max_concurrent_tasks", 2)),
        "task_runner.task_timeout_seconds": int(task_runner.get("task_timeout_seconds", 21600)),
        "daily_duplicate.local_worker_count": int(dd.get("local_worker_count", dd.get("report_worker_count", 3))),
        "daily_duplicate.report_worker_count": int(dd.get("report_worker_count", 3)),
        "daily_duplicate.llm_candidate_top_n": int(dd.get("llm_candidate_top_n", 3)),
        "daily_duplicate.llm_candidate_score_threshold": float(dd.get("llm_candidate_score_threshold", 0.72)),
        "daily_duplicate.low_info_candidate_score_threshold": float(dd.get("low_info_candidate_score_threshold", 0.82)),
        "daily_duplicate.max_candidates": int(dd.get("max_candidates", 20)),
        "daily_duplicate.self_history_days": int(dd.get("self_history_days", 30)),
        "daily_duplicate.cross_user_days": int(dd.get("cross_user_days", 3)),
        "daily_duplicate.cross_user_scope": str(dd.get("cross_user_scope", "department")),
        "daily_duplicate.score_weights.text": float(weights.get("text", 0.45)),
        "daily_duplicate.score_weights.semantic": float(weights.get("semantic", 0.45)),
        "daily_duplicate.score_weights.recency": float(weights.get("recency", 0.10)),
        "llm_judge.enabled": bool(llm.get("enabled", False)),
        "llm_judge.base_url": str(llm.get("base_url", "")),
        "llm_judge.chat_path": str(llm.get("chat_path", "/v1/chat/completions")),
        "llm_judge.model": str(llm.get("model", "")),
        "llm_judge.api_key_file": str(llm.get("api_key_file", "")),
        "llm_judge.api_key_env": str(llm.get("api_key_env", "DAILY_REPORT_LLM_API_KEY")),
        "llm_judge.global_max_concurrency": int(llm.get("global_max_concurrency", 30)),
        "llm_judge.per_task_max_concurrency": int(llm.get("per_task_max_concurrency", 30)),
        "llm_judge.request_timeout_seconds": int(llm.get("request_timeout_seconds", 180)),
        "llm_judge.max_retries": int(llm.get("max_retries", 1)),
        "testcase_verification.enabled": bool(tc.get("enabled", True)),
        "testcase_verification.self_history_days": int(tc.get("self_history_days", 7)),
        "testcase_verification.self_history_high_rate": float(tc.get("self_history_high_rate", 0.50)),
        "testcase_verification.self_history_medium_rate": float(tc.get("self_history_medium_rate", 0.20)),
        "checks.enable_code_check": bool(checks.get("enable_code_check", True)),
        "checks.enable_document_check": bool(checks.get("enable_document_check", True)),
        "upload.max_file_mb": int(config.get("upload", {}).get("max_file_mb", 50)),
    }


def effective_settings(config: dict, db) -> dict:
    """Return only currently supported keys, dropping retired persisted settings."""
    defaults = flatten_settings(config)
    stored = db.get_settings(defaults)
    return {key: stored.get(key, default_value) for key, default_value in defaults.items()}


def apply_flat_settings(config: dict, settings: dict) -> dict:
    merged = json.loads(json.dumps(config, ensure_ascii=False))
    for key, value in settings.items():
        target = merged
        parts = key.split(".")
        for part in parts[:-1]:
            target = target.setdefault(part, {})
        target[parts[-1]] = value
    return merged


def redact_audit_value(key: str, value):
    lowered = str(key).lower()
    if any(part in lowered for part in AUDIT_SENSITIVE_PARTS):
        return "[redacted]"
    if isinstance(value, dict):
        return {str(k): redact_audit_value(str(k), v) for k, v in value.items()}
    if isinstance(value, list):
        return [redact_audit_value(key, item) for item in value]
    return value


def audit_changes(before: dict | None, after: dict | None, allowed: tuple[str, ...]) -> dict:
    before, after = before or {}, after or {}
    changes = {}
    for key in allowed:
        if key not in after or before.get(key) == after.get(key):
            continue
        if key == "password":
            changes[key] = {"before": "[redacted]", "after": "[changed]"}
        else:
            changes[key] = {"before": redact_audit_value(key, before.get(key)), "after": redact_audit_value(key, after.get(key))}
    return changes


def csv_safe(value) -> str:
    text = json.dumps(value, ensure_ascii=False) if isinstance(value, (dict, list)) else str(value or "")
    if text.startswith(("=", "+", "-", "@", "\t", "\r")):
        return "'" + text
    return text


class DailyReportHandler(BaseHTTPRequestHandler):
    config: dict = {}
    db = None
    storage: FileStorage | None = None

    def _cookie_session_id(self) -> str:
        raw = self.headers.get("Cookie", "")
        for part in raw.split(";"):
            if "=" in part:
                k, v = part.strip().split("=", 1)
                if k == SESSION_COOKIE:
                    return v
        return ""

    def _current_user(self):
        return self.db.user_by_session(self._cookie_session_id())

    def _require_user(self):
        user = self._current_user()
        if not user:
            self._send_json({"ok": False, "error": "authentication_required"}, HTTPStatus.UNAUTHORIZED)
            return None
        return user

    def _require_manager(self):
        user = self._require_user()
        if user and not can_manage(user):
            self._send_json({"ok": False, "error": "forbidden"}, HTTPStatus.FORBIDDEN)
            return None
        return user

    def _require_business_user(self):
        user = self._require_user()
        if user and user.role not in BUSINESS_ROLES:
            self._send_json({"ok": False, "error": "forbidden"}, HTTPStatus.FORBIDDEN)
            return None
        return user

    def _require_report_exporter(self):
        user = self._require_business_user()
        if user and user.role == "employee":
            self._send_json({"ok": False, "error": "report_export_paused_for_employees"}, HTTPStatus.FORBIDDEN)
            return None
        return user

    def _require_admin(self):
        user = self._require_user()
        if user and not can_admin(user):
            self._send_json({"ok": False, "error": "forbidden"}, HTTPStatus.FORBIDDEN)
            return None
        return user

    @staticmethod
    def _audit_descriptor(method: str, path: str) -> dict | None:
        exact = {
            ("POST", "/api/auth/login"): ("authentication", "auth.login", "登录", "account", "info"),
            ("POST", "/api/auth/logout"): ("authentication", "auth.logout", "登出", "session", "info"),
            ("POST", "/api/auth/register"): ("authentication", "auth.register", "自助注册", "account", "info"),
            ("POST", "/api/testcase-snapshots"): ("business", "snapshot.upsert", "上传或替换测试总表", "testcase_snapshot", "info"),
            ("POST", "/api/run"): ("business", "task.create", "创建查重任务", "check_task", "info"),
            ("POST", "/api/tasks"): ("business", "task.create", "创建查重任务", "check_task", "info"),
            ("POST", "/api/admin/test-llm"): ("administration", "admin.llm.test", "测试 LLM 连接", "llm", "info"),
            ("POST", "/api/admin/cleanup-sessions"): ("administration", "admin.sessions.cleanup", "清理过期会话", "session", "info"),
            ("POST", "/api/admin/reset/prepare"): ("administration", "admin.reset.prepare", "准备清除业务数据", "business_data", "info"),
            ("POST", "/api/admin/reset/confirm"): ("administration", "admin.reset.confirm", "清除业务数据", "business_data", "critical"),
            ("POST", "/api/users"): ("administration", "admin.user.create", "新增账号", "user", "info"),
            ("POST", "/api/departments"): ("administration", "admin.department.create", "新增部门", "department", "info"),
            ("POST", "/api/groups"): ("administration", "admin.group.create", "新增小组", "group", "info"),
            ("POST", "/api/admin/settings"): ("administration", "admin.settings.update", "修改系统参数", "setting", "info"),
            ("PUT", "/api/admin/settings"): ("administration", "admin.settings.update", "修改系统参数", "setting", "info"),
            ("GET", "/api/export/reports.docx"): ("export", "export.weekly", "导出日报汇总", "report_export", "info"),
            ("GET", "/api/export/periodic.docx"): ("export", "export.periodic", "导出周期报表", "report_export", "info"),
            ("GET", "/api/export/periodic.xlsx"): ("export", "export.periodic", "导出周期报表", "report_export", "info"),
            ("GET", "/api/admin/audit-logs.csv"): ("export", "audit.export", "导出审查日志", "audit_log", "info"),
        }
        item = exact.get((method, path))
        if not item:
            if method == "POST" and path.startswith("/api/tasks/") and path.endswith("/retry"):
                item = ("business", "task.retry", "重试查重任务", "check_task", "info")
            elif method == "POST" and path.startswith("/api/users/") and path.endswith("/reset-password"):
                item = ("administration", "admin.user.password_reset", "重置账号密码", "user", "critical")
            elif method == "POST" and path.startswith("/api/tasks/") and path.endswith("/cancel"):
                item = ("business", "task.cancel", "取消查重任务", "check_task", "critical")
            elif method == "POST" and path.startswith("/api/admin/tasks/"):
                action_name = path.rsplit("/", 1)[-1]
                actions = {
                    "pause": ("task.pause", "暂停查重任务", "info"),
                    "resume": ("task.resume", "恢复查重任务", "info"),
                    "restart": ("task.restart", "重启查重任务", "warning"),
                    "force-terminate": ("task.force_terminate", "强制终止查重任务", "critical"),
                }
                if action_name in actions:
                    action, label, severity = actions[action_name]
                    item = ("administration", action, label, "check_task", severity)
            elif method == "PUT" and path.startswith("/api/my/reports/"):
                item = ("business", "report.update", "修改日报", "report", "info")
            elif method == "GET" and path.startswith("/api/tasks/") and (path.endswith("/export.xlsx") or path.endswith("/export.docx") or path.endswith("/missing-reports.xlsx")):
                item = ("export", "export.task_result", "导出查重结果", "check_task", "info")
            else:
                for prefix, action, label, target in (("/api/users/", "admin.user", "账号", "user"), ("/api/departments/", "admin.department", "部门", "department"), ("/api/groups/", "admin.group", "小组", "group"), ("/api/tasks/", "task", "查重任务", "check_task")):
                    if path.startswith(prefix) and method in {"PUT", "DELETE"}:
                        verb = "update" if method == "PUT" else "delete"
                        item = ("business" if prefix == "/api/tasks/" else "administration", f"{action}.{verb}", f"{'修改' if verb == 'update' else '删除'}{label}", target, "critical" if verb == "delete" else "info")
                        break
        if not item:
            protected_reads = {"/api/users", "/api/departments", "/api/groups", "/api/admin/settings", "/api/admin/storage-stats", "/api/admin/audit-logs", "/api/admin/runtime-logs", "/api/admin/runtime-logs/stream", "/api/admin/resource-monitor/snapshot", "/api/admin/resource-monitor/history"}
            if method == "GET" and path in protected_reads:
                return {"category": "administration", "action": "security.permission_denied", "summary": "访问管理接口",
                        "target_type": "admin_endpoint", "target_id": path, "target_label": path,
                        "severity": "warning", "changes": {}, "failure_only": True}
            return None
        category, action, summary, target_type, severity = item
        parts = [part for part in path.split("/") if part]
        target_id = ""
        if target_type in {"user", "department", "group", "check_task", "report"} and len(parts) >= 3:
            target_id = parts[3] if len(parts) > 4 and parts[1:3] == ["admin", "tasks"] else (parts[2] if parts[1] in {"users", "departments", "groups", "tasks"} else (parts[3] if len(parts) > 3 else ""))
        return {"category": category, "action": action, "summary": summary, "target_type": target_type,
                "target_id": target_id, "target_label": "", "severity": severity, "changes": {}}

    def _begin_audit(self, method: str, path: str) -> None:
        self._pending_audit = self._audit_descriptor(method, path)
        if not self._pending_audit:
            return
        user = self._current_user()
        if user:
            self._pending_audit.update({"actor_user_id": user.user_id, "actor_username": user.username,
                                        "actor_display_name": user.display_name, "actor_role": user.role})

    def _audit_update(self, **values) -> None:
        if getattr(self, "_pending_audit", None):
            self._pending_audit.update({key: value for key, value in values.items() if value is not None})

    def _finish_audit(self, status: int, response=None) -> None:
        pending = getattr(self, "_pending_audit", None)
        self._pending_audit = None
        if not pending:
            return
        success = int(status) < 400 and not (isinstance(response, dict) and response.get("ok") is False)
        if success and pending.pop("failure_only", False):
            return
        pending["outcome"] = "success" if success else "failure"
        pending["severity"] = pending.get("severity", "info") if success else "warning"
        error = response.get("error", "") if isinstance(response, dict) else ""
        pending["error_type"] = str(error or ("" if success else f"HTTP {int(status)}"))
        pending["summary"] = f"{pending.get('summary', '操作')}{'成功' if success else '失败'}"
        forwarded = str(self.headers.get("X-Forwarded-For", "")).split(",", 1)[0].strip()
        pending["client_ip"] = forwarded or str(self.headers.get("X-Real-IP", "")).strip() or str(self.client_address[0])
        pending["user_agent"] = str(self.headers.get("User-Agent", ""))
        try:
            self.db.append_audit_log(pending)
        except Exception as exc:
            runtime_log(self.db, "web", "audit_write_failed", "审查日志写入失败", level="error", context={"action": pending.get("action"), "outcome": pending.get("outcome")}, exc=exc)

    @staticmethod
    def _audit_filters(query: dict[str, list[str]]) -> dict[str, str]:
        filters = {key: str((query.get(key) or [""])[0]).strip() for key in ("start", "end", "category", "severity", "outcome", "actor", "q")}
        if filters["category"] and filters["category"] not in AUDIT_CATEGORIES:
            raise ValueError("invalid audit category")
        if filters["severity"] and filters["severity"] not in AUDIT_SEVERITIES:
            raise ValueError("invalid audit severity")
        if filters["outcome"] and filters["outcome"] not in AUDIT_OUTCOMES:
            raise ValueError("invalid audit outcome")
        for key in ("start", "end"):
            if filters[key]: parse_date(filters[key])
        return filters

    @staticmethod
    def _runtime_log_filters(query: dict[str, list[str]]) -> dict[str, str]:
        filters = {key: str((query.get(key) or [""])[0]).strip() for key in ("start", "end", "level", "source", "task_id", "q")}
        if filters["level"] and filters["level"] not in RUNTIME_LOG_LEVELS:
            raise ValueError("invalid runtime log level")
        if filters["source"] and filters["source"] not in RUNTIME_LOG_SOURCES:
            raise ValueError("invalid runtime log source")
        for key in ("start", "end"):
            if filters[key]: parse_date(filters[key])
        return filters

    def _stream_runtime_logs(self, filters: dict[str, str], after_seq: int) -> None:
        self.send_response(HTTPStatus.OK)
        self.send_header("Content-Type", "text/event-stream; charset=utf-8")
        self.send_header("Cache-Control", "no-cache, no-store")
        self.send_header("Connection", "keep-alive")
        self.send_header("X-Accel-Buffering", "no")
        self.end_headers()
        cursor = max(0, after_seq)
        until, last_heartbeat = time.monotonic() + 25, 0.0
        try:
            while time.monotonic() < until:
                rows = self.db.stream_runtime_logs(filters, cursor, 100)
                for row in rows:
                    cursor = max(cursor, int(row.get("seq") or 0))
                    self.wfile.write(f"id: {cursor}\nevent: runtime_log\ndata: {runtime_log_json(row)}\n\n".encode("utf-8"))
                now = time.monotonic()
                if rows or now - last_heartbeat >= 10:
                    if not rows:
                        self.wfile.write(b": keepalive\n\n")
                    self.wfile.flush()
                    last_heartbeat = now
                time.sleep(0.5)
        except (BrokenPipeError, ConnectionResetError, TimeoutError):
            return

    @staticmethod
    def _audit_csv(rows: list[dict]) -> bytes:
        output = io.StringIO(newline="")
        writer = csv.writer(output)
        writer.writerow(["时间", "级别", "类别", "结果", "操作者账号", "操作者姓名", "角色", "动作", "目标类型", "目标ID", "目标", "摘要", "脱敏差异", "错误", "来源IP", "User-Agent"])
        for row in rows:
            writer.writerow([csv_safe(row.get(key)) for key in ("occurred_at", "severity", "category", "outcome", "actor_username",
                "actor_display_name", "actor_role", "action", "target_type", "target_id", "target_label", "summary", "changes",
                "error_type", "client_ip", "user_agent")])
        return b"\xef\xbb\xbf" + output.getvalue().encode("utf-8")

    def do_GET(self) -> None:
        parsed = urlparse(self.path)
        path = unquote(parsed.path)
        query = parse_qs(parsed.query)
        self._begin_audit("GET", path)
        try:
            if path == "/api/time":
                self._send_json({"timezone": "Asia/Shanghai", "now": now_iso()})
            elif path == "/api/auth/me":
                user = self._current_user()
                self._send_json({"authenticated": bool(user), "user": user.to_dict() if user else None})
            elif path == "/api/public/departments":
                self._send_json(self.db.list_departments(public_only=True))
            elif path == "/api/public/groups":
                self._send_json(self.db.list_groups((query.get("department_id") or [""])[0], public_only=True))
            elif path == "/api/my/reports":
                user = self._require_business_user()
                if user:
                    filters = {k: (query.get(k) or [""])[0] for k in ("start", "end", "name", "role")}
                    reports = self.db.list_reports(user, filters["start"], filters["end"], True, filters["name"], filters["role"])
                    self._send_json([self._report_summary(r, user) for r in reports])
            elif path.startswith("/api/my/reports/") and path.endswith("/download"):
                user = self._require_business_user(); rid = path.split("/")[4]
                if user:
                    report = self.db.get_report(rid, user)
                    if not report: return self._send_json({"ok": False, "error": "not_found"}, HTTPStatus.NOT_FOUND)
                    self._send_bytes(self.storage.read(report.stored_path), "application/vnd.openxmlformats-officedocument.wordprocessingml.document", report.original_filename or "report.docx")
            elif path.startswith("/api/my/reports/"):
                user = self._require_business_user(); rid = path.split("/")[4]
                if user:
                    report = self.db.get_report(rid, user)
                    if not report or report.owner_user_id != user.user_id:
                        return self._send_json({"ok": False, "error": "not_found"}, HTTPStatus.NOT_FOUND)
                    self._send_json(self._report_detail(report))
            elif path.startswith("/api/artifacts/") and path.endswith("/download"):
                user = self._require_user(); aid = path.split("/")[3]
                if user:
                    art = self.db.get_artifact(aid, user)
                    if not art: return self._send_json({"ok": False, "error": "not_found"}, HTTPStatus.NOT_FOUND)
                    self._send_bytes(self.storage.read(art["stored_path"]), art.get("content_type") or "application/octet-stream", art.get("original_filename") or "artifact.bin")
            elif path == "/api/status":
                user = self._require_manager()
                if user:
                    tasks = self.db.list_tasks(user)
                    latest = self._task_for_view(tasks[0], user) if tasks else {}
                    s = latest.get("summary", {}) if latest else {}
                    report_count = len(self.db.list_reports(None) if user.role == "admin" else self.db.list_reports(user))
                    self._send_json({"latest_task": latest, "result_count": s.get("report_count", 0), "high_risk_count": s.get("high_risk_count", 0), "error_count": 1 if latest.get("status") == "failed" else 0, "metadata": {"report_count": report_count}})
            elif path == "/api/results":
                user = self._require_manager()
                if user:
                    tasks = self.db.list_tasks(user)
                    if not tasks: return self._send_json([])
                    task = self.db.get_task(tasks[0]["task_id"], user)
                    self._send_json(self._visible_task_cases(task or {}, user))
            elif path == "/api/tasks":
                user = self._require_manager()
                if user:
                    include_hidden = (query.get("include_hidden") or ["0"])[0] == "1"
                    tasks = self.db.list_tasks(user, {"status": (query.get("status") or [""])[0], "start": (query.get("start") or [""])[0], "end": (query.get("end") or [""])[0]}, include_hidden=include_hidden)
                    self._send_json([self._task_for_view(task, user) for task in tasks])
            elif path.startswith("/api/tasks/") and path.endswith("/status"):
                user = self._require_manager(); tid = path.split("/")[3]
                if user:
                    task = self.db.get_task_status(tid, user)
                    if not task: return self._send_json({"ok": False, "error": "not_found"}, HTTPStatus.NOT_FOUND)
                    self._send_json(self._task_for_view(task, user))
            elif path.startswith("/api/tasks/") and path.endswith("/results"):
                user = self._require_manager(); tid = path.split("/")[3]
                if user:
                    task = self.db.get_task(tid, user)
                    if not task: return self._send_json({"ok": False, "error": "not_found"}, HTTPStatus.NOT_FOUND)
                    self._send_json(self._visible_task_cases(task, user))
            elif path.startswith("/api/tasks/") and path.endswith("/missing-reports"):
                user = self._require_manager(); tid = path.split("/")[3]
                if user:
                    task = self.db.get_task(tid, user)
                    if not task: return self._send_json({"ok": False, "error": "not_found"}, HTTPStatus.NOT_FOUND)
                    self._send_json(self._visible_missing_reports(task, user))
            elif path.startswith("/api/tasks/") and path.endswith("/export.xlsx"):
                self._send_task_file(path.split("/")[3], "review_cases.xlsx", "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
            elif path.startswith("/api/tasks/") and path.endswith("/export.docx"):
                self._send_task_file(path.split("/")[3], "review_cases.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            elif path.startswith("/api/tasks/") and path.endswith("/missing-reports.xlsx"):
                self._send_task_file(path.split("/")[3], "missing_reports.xlsx", "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
            elif path == "/api/testcase-snapshots":
                user = self._require_user()
                if user: self._send_json(self.db.list_testcase_snapshots(user))
            elif path.startswith("/api/testcase-snapshots/") and path.endswith("/download"):
                user = self._require_user(); sid = path.split("/")[3]
                if user:
                    snap = self.db.get_testcase_snapshot(sid, user)
                    if not snap: return self._send_json({"ok": False, "error": "not_found"}, HTTPStatus.NOT_FOUND)
                    self._send_bytes(self.storage.read(snap["stored_path"]), "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", snap["original_filename"])
            elif path == "/api/export/reports.docx":
                user = self._require_report_exporter()
                if user:
                    start, end = (query.get("start") or [""])[0], (query.get("end") or [""])[0]
                    name, role = (query.get("name") or [""])[0], (query.get("role") or [""])[0]
                    reports = self.db.list_reports(user, start, end, True, name, role)
                    body = _build_reports_docx(reports)
                    stored = self.storage.save_bytes("exports/weekly_reports", f"weekly_{start or 'all'}_{end or 'all'}.docx", body, end or start or None)
                    self.db.save_export({"export_type": "weekly_reports", "created_by": user.user_id, "start_date": start, "end_date": end, "stored_path": stored.stored_path})
                    self._audit_update(target_label=f"{len(reports)} 份日报", changes={"start": start, "end": end, "report_count": len(reports)})
                    self._send_bytes(body, "application/vnd.openxmlformats-officedocument.wordprocessingml.document", stored.original_filename)
            elif path in {"/api/export/periodic.docx", "/api/export/periodic.xlsx"}:
                user = self._require_report_exporter()
                if user:
                    report_type, start, end = _periodic_export_args(query)
                    reports = self.db.list_reports(user, start, end)
                    artifacts = {report.report_id: self.db.list_artifacts_for_report(report.report_id) for report in reports}
                    missing_reports = _periodic_missing_reports(self.db, user, start, end, reports)
                    title = f"{'周报' if report_type == 'weekly' else '月报'}汇总"
                    scope = _export_scope_label(user)
                    if path.endswith(".docx"):
                        body = _build_periodic_reports_docx(title, scope, start, end, reports, artifacts, missing_reports)
                        filename = f"{report_type}_{start}_{end}_{scope}.docx"
                        content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    else:
                        body = _build_periodic_reports_xlsx(title, scope, start, end, reports, artifacts, missing_reports)
                        filename = f"{report_type}_{start}_{end}_{scope}.xlsx"
                        content_type = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                    stored = self.storage.save_bytes("exports/periodic_reports", filename, body, end)
                    self.db.save_export({"export_type": f"{report_type}_{'docx' if path.endswith('.docx') else 'xlsx'}", "created_by": user.user_id, "department_scope": user.department_id, "start_date": start, "end_date": end, "stored_path": stored.stored_path})
                    self._audit_update(target_label=f"{len(reports)} 份日报", changes={"report_type": report_type, "start": start, "end": end, "report_count": len(reports)})
                    self._send_bytes(body, content_type, stored.original_filename)
            elif path == "/api/users":
                user = self._require_admin();
                if user: self._send_json(self.db.list_users())
            elif path == "/api/departments":
                user = self._require_admin();
                if user: self._send_json(self.db.list_departments())
            elif path == "/api/groups":
                user = self._require_admin()
                if user: self._send_json(self.db.list_groups((query.get("department_id") or [""])[0]))
            elif path == "/api/admin/settings":
                user = self._require_admin();
                if user: self._send_json(effective_settings(self.config, self.db))
            elif path == "/api/admin/task-console":
                user = self._require_admin()
                if user:
                    try: limit = int((query.get("limit") or ["50"])[0])
                    except ValueError: raise ValueError("invalid task page limit")
                    filters = {
                        "view": (query.get("view") or ["active"])[0],
                        "status": (query.get("status") or [""])[0],
                        "creator": (query.get("creator") or [""])[0],
                        "trigger_source": (query.get("trigger_source") or [""])[0],
                        "q": (query.get("q") or [""])[0],
                    }
                    data = self.db.list_admin_tasks(filters, str((query.get("cursor") or [""])[0]), limit)
                    self._send_json(self._task_console_response(data))
            elif path == "/api/admin/audit-logs":
                user = self._require_admin()
                if user:
                    filters = self._audit_filters(query)
                    cursor = str((query.get("cursor") or [""])[0])
                    try: limit = int((query.get("limit") or ["50"])[0])
                    except ValueError: raise ValueError("invalid audit page limit")
                    self._send_json(self.db.list_audit_logs(filters, cursor, limit))
            elif path == "/api/admin/audit-logs.csv":
                user = self._require_admin()
                if user:
                    filters = self._audit_filters(query)
                    rows = self.db.export_audit_logs(filters)
                    self._audit_update(target_label=f"{len(rows)} 条", changes={"filters": redact_audit_value("filters", filters)})
                    self._send_bytes(self._audit_csv(rows), "text/csv; charset=utf-8", f"audit_logs_{now_iso()[:10]}.csv")
            elif path == "/api/admin/runtime-logs/stream":
                user = self._require_admin()
                if user:
                    filters = self._runtime_log_filters(query)
                    raw_after = str(self.headers.get("Last-Event-ID") or (query.get("after") or ["0"])[0])
                    try: after_seq = max(0, int(raw_after))
                    except ValueError: raise ValueError("invalid runtime log sequence")
                    self._stream_runtime_logs(filters, after_seq)
            elif path == "/api/admin/runtime-logs":
                user = self._require_admin()
                if user:
                    filters = self._runtime_log_filters(query)
                    cursor = str((query.get("cursor") or [""])[0])
                    try: limit = int((query.get("limit") or ["100"])[0])
                    except ValueError: raise ValueError("invalid runtime log page limit")
                    self._send_json(self.db.list_runtime_logs(filters, cursor, limit))
            elif path == "/api/admin/storage-stats":
                user = self._require_admin();
                if user: self._send_json(self._storage_stats())
            elif path == "/api/admin/resource-monitor/snapshot":
                user = self._require_admin()
                if user: self._send_json(build_resource_snapshot(self.config, self.db))
            elif path == "/api/admin/resource-monitor/history":
                user = self._require_admin()
                if user: self._send_json(build_resource_history(self.config, str((query.get("range") or ["1h"])[0])))
            elif path == "/" or path == "/index.html":
                self._send_file(Path(self.config["app"]["frontend_dir"]) / "index.html")
            else:
                frontend = Path(self.config["app"]["frontend_dir"]).resolve()
                static_path = (frontend / path.lstrip("/")).resolve()
                if frontend in static_path.parents and static_path.exists() and static_path.is_file(): self._send_file(static_path)
                else: self.send_error(HTTPStatus.NOT_FOUND, "Not found")
        except Exception as exc:
            runtime_log(self.db, "web", "request_failed", "处理 GET 请求时发生异常", level="error", context={"path": path}, exc=exc)
            self._send_json({"ok": False, "error": str(exc)}, HTTPStatus.INTERNAL_SERVER_ERROR)

    def do_POST(self) -> None:
        path = unquote(urlparse(self.path).path)
        self._begin_audit("POST", path)
        try:
            if path == "/api/auth/login":
                payload = self._read_json_body(); username = str(payload.get("username") or ""); self._audit_update(actor_username=username, target_label=username)
                user = self.db.authenticate(username, str(payload.get("password") or ""))
                if not user: return self._send_json({"ok": False, "error": "invalid_credentials"}, HTTPStatus.UNAUTHORIZED)
                self._audit_update(actor_user_id=user.user_id, actor_username=user.username, actor_display_name=user.display_name, actor_role=user.role, target_id=user.user_id, target_label=user.username)
                sid = self.db.create_session(user.user_id); secure = "; Secure" if self.config.get("app", {}).get("secure_cookie") else ""
                self._send_json({"ok": True, "user": user.to_dict()}, headers={"Set-Cookie": f"{SESSION_COOKIE}={sid}; Path=/; HttpOnly; SameSite=Lax{secure}"})
            elif path == "/api/auth/logout":
                self.db.delete_session(self._cookie_session_id()); self._send_json({"ok": True}, headers={"Set-Cookie": f"{SESSION_COOKIE}=; Path=/; Max-Age=0; HttpOnly; SameSite=Lax"})
            elif path == "/api/auth/register":
                if not self.config.get("app", {}).get("allow_self_register", True): return self._send_json({"ok": False, "error": "registration_disabled"}, HTTPStatus.FORBIDDEN)
                payload = self._read_json_body(); payload["role"] = "employee"; self._audit_update(actor_username=str(payload.get("username") or ""), target_label=str(payload.get("username") or ""))
                created = self.db.create_user(payload); self._audit_update(actor_user_id=created.get("user_id", ""), target_id=created.get("user_id", ""), changes=audit_changes({}, payload, ("username", "display_name", "role", "department_id", "group_id")))
                self._send_json({"ok": True, **created}, HTTPStatus.CREATED)
            elif path == "/api/reports":
                user = self._require_business_user()
                if user:
                    fields, files = self._read_multipart_or_json()
                    payload = json.loads(fields.get("payload", "{}")) if "payload" in fields else fields
                    testcase_files = files.get("testcase_files", [])
                    if testcase_files:
                        payload["testcase_items"] = self._parse_testcase_upload(testcase_files[0])
                    report = self._report_from_payload(user, payload)
                    self.db.save_report(report)
                    for f in files.get("code_files", []): self._save_artifact(report, user, f, "code")
                    for f in files.get("document_files", []): self._save_artifact(report, user, f, "document")
                    for f in testcase_files: self._save_artifact(report, user, f, "testcase")
                    self._send_json({"ok": True, "report_id": report.report_id, "file": report.original_filename}, HTTPStatus.CREATED)
            elif path == "/api/testcase-snapshots":
                user = self._require_manager()
                if user:
                    if user.role not in {ROLE_GROUP_LEADER, ROLE_DIRECTOR}:
                        return self._send_json({"ok": False, "error": "only_group_leader_or_director_uploads_snapshot"}, HTTPStatus.FORBIDDEN)
                    fields, files = self._read_multipart_or_json()
                    snapshot_date = parse_date(fields.get("snapshot_date", ""))
                    file_items = files.get("snapshot_file", [])
                    if not file_items: raise ValueError("missing snapshot_file")
                    f = file_items[0]
                    group_id, group_name, department_id, department_name = self._resolve_snapshot_group(user, fields)
                    # Validate sheet before saving.
                    from src.core.testcase_excel import load_daily_sheet_rows
                    load_daily_sheet_rows(f.data)
                    stored = self.storage.save_bytes(f"testcase_snapshots/{group_id}", f.filename, f.data, snapshot_date)
                    sid = self.db.save_testcase_snapshot({
                        "department_id": department_id,
                        "department_name": department_name,
                        "group_id": group_id,
                        "group_name": group_name,
                        "snapshot_date": snapshot_date,
                        "original_filename": f.filename,
                        "stored_path": stored.stored_path,
                        "file_size": stored.file_size,
                        "content_hash": stored.content_hash,
                        "uploaded_by": user.user_id,
                    })
                    self._audit_update(target_id=sid, target_label=f"{group_name} {snapshot_date}", changes={"group_id": group_id, "snapshot_date": snapshot_date, "filename": Path(f.filename).name})
                    self._send_json({"ok": True, "snapshot_id": sid}, HTTPStatus.CREATED)
            elif path in {"/api/run", "/api/tasks"}:
                user = self._require_manager()
                if user:
                    report_date = self._new_report_date()
                    payload = {"report_date_start": report_date, "report_date_end": report_date} if path == "/api/run" else self._read_json_body()
                    start, end = str(payload.get("report_date_start") or ""), str(payload.get("report_date_end") or "")
                    if not start or not end:
                        raise ValueError("查重任务必须选择开始日期和结束日期")
                    if parse_date(start) > parse_date(end):
                        raise ValueError("开始日期不能晚于结束日期")
                    cfg = apply_flat_settings(self.config, effective_settings(self.config, self.db))
                    rel_dir = f"tasks/pending_{token('', 8)}"
                    task_id = self.db.create_task(user, payload, cfg, rel_dir)
                    task = self.db.get_task(task_id, user)
                    self._audit_update(target_id=task_id, target_label=(task or {}).get("task_name", task_id), changes={"report_date_start": start, "report_date_end": end})
                    self._send_json({
                        "ok": True,
                        "task_id": task_id,
                        "task_name": (task or {}).get("task_name", task_id),
                        "status": "queued",
                    }, HTTPStatus.ACCEPTED)
            elif path.startswith("/api/tasks/") and path.endswith("/retry"):
                user = self._require_manager(); task_id = path.split("/")[3]
                if user:
                    self._audit_update(target_id=task_id)
                    if not self.db.retry_task(task_id, user):
                        return self._send_json({"ok": False, "error": "task_not_retryable"}, HTTPStatus.CONFLICT)
                    self._send_json({"ok": True, "task_id": task_id, "status": "queued"}, HTTPStatus.ACCEPTED)
            elif path.startswith("/api/tasks/") and path.endswith("/cancel"):
                user = self._require_manager(); task_id = path.split("/")[3]
                if user:
                    self._audit_update(target_id=task_id)
                    task = self.db.get_task_status(task_id, user)
                    if not task: return self._send_json({"ok": False, "error": "not_found"}, HTTPStatus.NOT_FOUND)
                    if task.get("trigger_source") == "system_daily" and not can_admin(user):
                        return self._send_json({"ok": False, "error": "only_admin_can_cancel_system_task"}, HTTPStatus.FORBIDDEN)
                    status = self.db.cancel_task(task_id, user)
                    self._send_json({"ok": True, "task_id": task_id, "status": status}, HTTPStatus.ACCEPTED)
            elif path.startswith("/api/admin/tasks/") and path.rsplit("/", 1)[-1] in {"pause", "resume", "restart", "force-terminate"}:
                user = self._require_admin()
                if user:
                    parts = path.split("/"); task_id, action = parts[4], parts[5]
                    task = self.db.get_task(task_id)
                    if not task:
                        return self._send_json({"ok": False, "error": "task_not_found"}, HTTPStatus.NOT_FOUND)
                    payload = self._read_json_body()
                    if action == "force-terminate" and str(payload.get("confirmation_task_id") or "") != task_id:
                        return self._send_json({"ok": False, "error": "confirmation_task_id_mismatch"}, HTTPStatus.BAD_REQUEST)
                    before = str(task.get("effective_status") or task.get("status") or "")
                    if action == "pause": status = self.db.admin_pause_task(task_id)
                    elif action == "resume": status = self.db.admin_resume_task(task_id)
                    elif action == "restart": status = self.db.admin_restart_task(task_id)
                    else: status = self.db.admin_force_terminate_task(task_id)
                    updated = self.db.get_task(task_id) or task
                    self._audit_update(target_id=task_id, target_label=str(task.get("task_name") or task_id), changes={
                        "from_status": before, "to_status": status, "attempt": int(updated.get("current_attempt") or 1),
                    })
                    self._send_json({"ok": True, "task_id": task_id, "status": status, "attempt": int(updated.get("current_attempt") or 1)}, HTTPStatus.ACCEPTED)
            elif path == "/api/admin/test-llm":
                user = self._require_admin()
                if user:
                    cfg = apply_flat_settings(self.config, effective_settings(self.config, self.db))
                    judge = create_llm_judge(cfg)
                    self._send_json(judge.health_check() if judge else {"ok": False, "error": "llm disabled"})
            elif path == "/api/admin/cleanup-sessions":
                user = self._require_admin();
                if user:
                    deleted = self.db.cleanup_sessions(); self._audit_update(changes={"deleted_count": deleted}); self._send_json({"ok": True, "deleted": deleted})
            elif path == "/api/tasks/display/clear":
                user = self._require_manager()
                if user: self.db.clear_task_page_display(user); self._send_json({"ok": True})
            elif path == "/api/admin/reset/prepare":
                user = self._require_admin()
                if user:
                    self._send_json({"ok": True, "confirmation_id": self.db.create_reset_confirmation(user.user_id), "confirmation_phrase": RESET_CONFIRMATION_PHRASE})
            elif path == "/api/admin/reset/confirm":
                user = self._require_admin()
                if user:
                    payload = self._read_json_body()
                    if str(payload.get("confirmation_phrase") or "") != RESET_CONFIRMATION_PHRASE:
                        raise ValueError("confirmation phrase does not match")
                    if not self.db.consume_reset_confirmation(str(payload.get("confirmation_id") or ""), user.user_id):
                        raise ValueError("confirmation expired or invalid")
                    counts = self.db.clear_business_data()
                    removed = self._clear_business_storage()
                    self._audit_update(changes={"deleted_rows": counts, "removed_files": removed})
                    self._send_json({"ok": True, "deleted": counts, "removed_files": removed}, headers={"Set-Cookie": f"{SESSION_COOKIE}=; Path=/; Max-Age=0; HttpOnly; SameSite=Lax"})
            elif path == "/api/users":
                user = self._require_admin()
                if user:
                    payload = self._read_json_body(); created = self.db.create_user(payload)
                    self._audit_update(target_id=created.get("user_id", ""), target_label=str(payload.get("username") or ""), changes=audit_changes({}, payload, ("username", "display_name", "role", "department_id", "group_id", "enabled", "password")))
                    self._send_json({"ok": True, **created}, HTTPStatus.CREATED)
            elif path.startswith("/api/users/") and path.endswith("/reset-password"):
                user = self._require_admin()
                if user:
                    user_id = path.split("/")[3]
                    target = self.db.get_user_by_id(user_id)
                    if not target:
                        raise ValueError("user not found")
                    password = str(self._read_json_body().get("password") or "")
                    self.db.reset_user_password(user_id, password)
                    self._audit_update(target_id=user_id, target_label=target.username, changes=audit_changes({}, {"password": password}, ("password",)), severity="critical")
                    self._send_json({"ok": True})
            elif path == "/api/departments":
                user = self._require_admin()
                if user:
                    payload = self._read_json_body(); created = self.db.create_department(payload)
                    self._audit_update(target_id=created.get("department_id", ""), target_label=created.get("name", ""), changes=audit_changes({}, payload, ("name", "parent_department_id", "enabled")))
                    self._send_json({"ok": True, "department": created}, HTTPStatus.CREATED)
            elif path == "/api/groups":
                user = self._require_admin()
                if user:
                    payload = self._read_json_body(); created = self.db.create_group(payload)
                    self._audit_update(target_id=created.get("group_id", ""), target_label=created.get("name", ""), changes=audit_changes({}, payload, ("name", "department_id", "enabled")))
                    self._send_json({"ok": True, "group": created}, HTTPStatus.CREATED)
            elif path == "/api/admin/settings":
                user = self._require_admin()
                if user:
                    payload = self._read_json_body()
                    self._validate_settings_payload(payload)
                    before = effective_settings(self.config, self.db)
                    self.db.update_settings(payload)
                    after = before | payload
                    self._audit_update(changes=audit_changes(before, after, tuple(payload)))
                    self._send_json({"ok": True})
            else:
                self.send_error(HTTPStatus.NOT_FOUND, "Not found")
        except ValueError as exc:
            message = str(exc)
            conflict = message.startswith("active_task_exists:") or message.startswith("task_not_") or message.startswith("task_already_")
            status = HTTPStatus.REQUEST_ENTITY_TOO_LARGE if isinstance(exc, RequestTooLarge) else HTTPStatus.CONFLICT if conflict else HTTPStatus.BAD_REQUEST
            self._send_json({"ok": False, "error": message}, status)
        except Exception as exc:
            runtime_log(self.db, "web", "request_failed", "处理 POST 请求时发生异常", level="error", context={"path": path}, exc=exc)
            self._send_json({"ok": False, "error": str(exc)}, HTTPStatus.INTERNAL_SERVER_ERROR)

    def do_PUT(self) -> None:
        path = unquote(urlparse(self.path).path)
        self._begin_audit("PUT", path)
        try:
            if path.startswith("/api/users/"):
                user = self._require_admin()
                if user:
                    user_id = path.split("/")[3]; previous = self.db.get_user_by_id(user_id); before = previous.to_dict() if previous else {}
                    payload = self._read_json_body(); self.db.update_user(user_id, payload)
                    changes = audit_changes(before, payload, ("username", "display_name", "role", "department_id", "group_id", "manager_user_id", "enabled", "password"))
                    self._audit_update(target_id=user_id, target_label=before.get("username", ""), changes=changes, severity="critical" if {"role", "enabled", "password"} & changes.keys() else "info")
                    self._send_json({"ok": True})
            elif path.startswith("/api/departments/"):
                user = self._require_admin()
                if user:
                    department_id = path.split("/")[3]; before = next((row for row in self.db.list_departments() if row.get("department_id") == department_id), {})
                    payload = self._read_json_body(); self.db.update_department(department_id, payload)
                    self._audit_update(target_id=department_id, target_label=before.get("name", ""), changes=audit_changes(before, payload, ("name", "parent_department_id", "enabled")))
                    self._send_json({"ok": True})
            elif path.startswith("/api/groups/"):
                user = self._require_admin()
                if user:
                    group_id = path.split("/")[3]; before = next((row for row in self.db.list_groups() if row.get("group_id") == group_id), {})
                    payload = self._read_json_body(); self.db.update_group(group_id, payload)
                    self._audit_update(target_id=group_id, target_label=before.get("name", ""), changes=audit_changes(before, payload, ("name", "department_id", "enabled")))
                    self._send_json({"ok": True})
            elif path.startswith("/api/my/reports/"):
                user = self._require_business_user(); rid = path.split("/")[4]
                if user:
                    old = self.db.get_report(rid, user)
                    if not old or old.owner_user_id != user.user_id: raise ValueError("日报不存在或无权更改")
                    fields, files = self._read_multipart_or_json()
                    payload = json.loads(fields.get("payload", "{}")) if "payload" in fields else fields
                    testcase_files = files.get("testcase_files", [])
                    if testcase_files:
                        payload["testcase_items"] = self._parse_testcase_upload(testcase_files[0])
                    self._ensure_report_editable(old)
                    report = self._report_from_payload(user, payload, old.report_date, old.created_at)
                    report.report_id, report.updated_at = rid, now_iso()
                    self.db.save_report(report)
                    for f in files.get("code_files", []): self._save_artifact(report, user, f, "code")
                    for f in files.get("document_files", []): self._save_artifact(report, user, f, "document")
                    for f in testcase_files: self._save_artifact(report, user, f, "testcase")
                    changed_fields = [key for key in ("title_summary", "work_description", "code_items", "document_items", "testcase_items") if key in payload]
                    if files.get("code_files"): changed_fields.append("code_attachments")
                    if files.get("document_files"): changed_fields.append("document_attachments")
                    if testcase_files: changed_fields.append("testcase_attachment")
                    self._audit_update(target_id=rid, target_label=old.report_date, changes={"report_date": old.report_date, "changed_fields": changed_fields})
                    self._send_json({"ok": True})
            elif path == "/api/admin/settings":
                user = self._require_admin()
                if user:
                    payload = self._read_json_body()
                    self._validate_settings_payload(payload)
                    before = effective_settings(self.config, self.db)
                    self.db.update_settings(payload)
                    self._audit_update(changes=audit_changes(before, before | payload, tuple(payload)))
                    self._send_json({"ok": True})
            else: self.send_error(HTTPStatus.NOT_FOUND, "Not found")
        except ValueError as exc:
            status = HTTPStatus.REQUEST_ENTITY_TOO_LARGE if isinstance(exc, RequestTooLarge) else HTTPStatus.BAD_REQUEST
            self._send_json({"ok": False, "error": str(exc)}, status)
        except Exception as exc:
            runtime_log(self.db, "web", "request_failed", "处理 PUT 请求时发生异常", level="error", context={"path": path}, exc=exc)
            self._send_json({"ok": False, "error": str(exc)}, HTTPStatus.INTERNAL_SERVER_ERROR)

    def do_DELETE(self) -> None:
        path = unquote(urlparse(self.path).path)
        self._begin_audit("DELETE", path)
        try:
            if path.startswith("/api/users/"):
                user = self._require_admin()
                if user:
                    user_id = path.split("/")[3]; previous = self.db.get_user_by_id(user_id)
                    self._audit_update(target_id=user_id, target_label=previous.username if previous else "")
                    self.db.delete_user(user_id); self._send_json({"ok": True})
            elif path.startswith("/api/departments/"):
                user = self._require_admin()
                if user:
                    department_id = path.split("/")[3]; before = next((row for row in self.db.list_departments() if row.get("department_id") == department_id), {})
                    self._audit_update(target_id=department_id, target_label=before.get("name", "")); self.db.delete_department(department_id); self._send_json({"ok": True})
            elif path.startswith("/api/groups/"):
                user = self._require_admin()
                if user:
                    group_id = path.split("/")[3]; before = next((row for row in self.db.list_groups() if row.get("group_id") == group_id), {})
                    self._audit_update(target_id=group_id, target_label=before.get("name", "")); self.db.delete_group(group_id); self._send_json({"ok": True})
            elif path.startswith("/api/tasks/"):
                user = self._require_manager();
                if user:
                    task_id = path.split("/")[3]
                    task = self.db.get_task(task_id, user); self._audit_update(target_id=task_id, target_label=(task or {}).get("task_name", ""))
                    status = self.db.delete_task(task_id, user)
                    task_dir = (Path(self.config["storage"]["root_dir"]) / "tasks" / task_id).resolve()
                    storage_root = Path(self.config["storage"]["root_dir"]).resolve()
                    if status == "deleted" and storage_root in task_dir.parents and task_dir.exists(): shutil.rmtree(task_dir)
                    self._send_json({"ok": True, "status": status}, HTTPStatus.ACCEPTED if status == "deleting" else HTTPStatus.OK)
            else: self.send_error(HTTPStatus.NOT_FOUND, "Not found")
        except Exception as exc:
            runtime_log(self.db, "web", "request_failed", "处理 DELETE 请求时发生异常", level="error", context={"path": path}, exc=exc)
            self._send_json({"ok": False, "error": str(exc)}, HTTPStatus.INTERNAL_SERVER_ERROR)

    def _send_task_file(self, task_id: str, name: str, content_type: str) -> None:
        user = self._require_manager()
        if not user: return
        task = self.db.get_task(task_id, user)
        if not task: return self._send_json({"ok": False, "error": "not_found"}, HTTPStatus.NOT_FOUND)
        if task.get("trigger_source") != "system_daily":
            attempt = int(task.get("current_attempt") or 1)
            path = Path(self.config["storage"]["root_dir"]) / "tasks" / task_id / f"attempt_{attempt}" / name
            legacy_path = Path(self.config["storage"]["root_dir"]) / "tasks" / task_id / name
            if not path.exists() and legacy_path.exists(): path = legacy_path
            if not path.exists(): return self._send_json({"ok": False, "error": "file_not_found"}, HTTPStatus.NOT_FOUND)
            return self._send_bytes(path.read_bytes(), content_type, name)
        with tempfile.TemporaryDirectory() as temp_dir:
            output = Path(temp_dir) / name
            if name == "review_cases.xlsx":
                DailyReportPipeline.export_review_xlsx(self._visible_task_cases(task, user), output)
            elif name == "review_cases.docx":
                DailyReportPipeline.export_review_docx(self._visible_task_cases(task, user), output)
            elif name == "missing_reports.xlsx":
                DailyReportPipeline.export_missing_reports_xlsx(self._visible_missing_reports(task, user), output)
            else:
                return self._send_json({"ok": False, "error": "file_not_found"}, HTTPStatus.NOT_FOUND)
            self._send_bytes(output.read_bytes(), content_type, name)

    @staticmethod
    def _row_visible_to_manager(row: dict, user) -> bool:
        if user.role in {"minister", "admin"}:
            return True
        if user.role == "director":
            return row.get("department_id") == user.department_id or row.get("department_name") == user.department_name
        if user.role == "group_leader":
            return row.get("owner_user_id") == user.user_id or row.get("user_id") == user.user_id or row.get("group_id") == user.group_id or row.get("group_name") == user.group_name
        return False

    def _visible_task_cases(self, task: dict, user) -> list[dict]:
        cases = task.get("result", {}).get("report_cases", [])
        if task.get("trigger_source") != "system_daily":
            return cases
        return [case for case in cases if self._row_visible_to_manager(case, user)]

    def _visible_missing_reports(self, task: dict, user) -> list[dict]:
        rows = task.get("result", {}).get("missing_reports", [])
        if task.get("trigger_source") != "system_daily":
            return rows
        return [row for row in rows if self._row_visible_to_manager(row, user)]

    def _task_for_view(self, task: dict, user) -> dict:
        if task.get("trigger_source") != "system_daily":
            return task
        summary_for_user = getattr(self.db, "get_task_summary_for_user", None)
        if callable(summary_for_user) and task.get("partial_available"):
            view = dict(task)
            view["summary"] = summary_for_user(task["task_id"], int(task.get("current_attempt") or 1), user)
            return view
        full_task = self.db.get_task(task["task_id"], user) or task
        cases = self._visible_task_cases(full_task, user)
        missing = self._visible_missing_reports(full_task, user)
        view = dict(task)
        view["summary"] = {
            "report_count": len(cases),
            "high_risk_count": sum(1 for case in cases if case.get("overall_risk") == "high"),
            "medium_risk_count": sum(1 for case in cases if case.get("overall_risk") == "medium"),
            "missing_report_count": len(missing),
            **duplicate_usage_summary(cases),
        }
        return view

    def _task_console_response(self, data: dict) -> dict:
        tasks = list(data.get("tasks") or [])
        for task in tasks:
            status = str(task.get("status") or "")
            effective = str(task.get("effective_status") or status)
            actions: list[str] = []
            if status in {"queued", "running"} and effective not in {"pausing", "cancelling", "terminating", "deleting"}:
                actions.append("pause")
            if status == "paused":
                actions.extend(["resume", "restart"])
            elif status in {"failed", "cancelled", "terminated"}:
                actions.append("restart")
            if status in {"queued", "running", "paused"} and effective not in {"terminating", "deleting"}:
                actions.append("force_terminate")
            task["allowed_actions"] = actions
        lease = self.db.get_worker_lease()
        stale_seconds = max(30, int(self.config.get("task_runner", {}).get("stale_heartbeat_seconds", 90)))
        online = False
        if lease and lease.get("heartbeat_at"):
            try:
                online = datetime.fromisoformat(str(lease["heartbeat_at"])) >= datetime.fromisoformat(now_iso()) - timedelta(seconds=stale_seconds)
            except ValueError:
                online = False
        raw_counts = dict(data.get("counts") or {})
        counts = {
            "running": int(raw_counts.get("running", 0)),
            "queued": int(raw_counts.get("queued", 0)),
            "paused": int(raw_counts.get("paused", 0)),
            "abnormal": int(raw_counts.get("failed", 0)) + int(raw_counts.get("terminated", 0)),
        }
        return {
            "worker": {
                "online": online,
                "worker_id": str((lease or {}).get("worker_id") or ""),
                "heartbeat_at": str((lease or {}).get("heartbeat_at") or ""),
            },
            "counts": counts,
            "tasks": tasks,
            "next_cursor": str(data.get("next_cursor") or ""),
            "refreshed_at": now_iso(),
        }

    @staticmethod
    def _validate_settings_payload(payload: dict) -> None:
        for key, label in (("daily_report_submission.rollover_time", "日报归属切换时间"), ("automatic_daily_duplicate.run_at", "自动查重任务时间")):
            value = payload.get(key)
            if value is None:
                continue
            try:
                hour, minute = (int(part) for part in str(value).split(":", 1))
            except (TypeError, ValueError):
                raise ValueError(f"{label}须为 HH:MM 格式")
            if not (0 <= hour <= 23 and 0 <= minute <= 59):
                raise ValueError(f"{label}须在 00:00 至 23:59 之间")
        automatic_enabled = payload.get("automatic_daily_duplicate.enabled")
        if automatic_enabled is not None and not isinstance(automatic_enabled, bool):
            raise ValueError("自动查重任务开关必须为 true 或 false")
        task_limits = {
            "task_runner.max_concurrent_tasks": (1, 8, "任务并发数须为 1 至 8 的整数"),
            "task_runner.task_timeout_seconds": (600, 86400, "任务超时须为 600 至 86400 秒的整数"),
        }
        for key, (minimum, maximum, message) in task_limits.items():
            value = payload.get(key)
            if value is None:
                continue
            try:
                number = int(value)
            except (TypeError, ValueError):
                raise ValueError(message)
            if str(value).strip() != str(number) or not minimum <= number <= maximum:
                raise ValueError(message)
        worker_count = payload.get("daily_duplicate.report_worker_count")
        if worker_count is not None:
            try:
                workers = int(worker_count)
            except (TypeError, ValueError):
                raise ValueError("单任务并行数须为 1 至 8 的整数")
            if str(worker_count).strip() != str(workers) or not 1 <= workers <= 8:
                raise ValueError("单任务并行数须为 1 至 8 的整数")
        local_worker_count = payload.get("daily_duplicate.local_worker_count")
        if local_worker_count is not None:
            try:
                workers = int(local_worker_count)
            except (TypeError, ValueError):
                raise ValueError("本地查重并行数须为 1 至 32 的整数")
            if str(local_worker_count).strip() != str(workers) or not 1 <= workers <= 32:
                raise ValueError("本地查重并行数须为 1 至 32 的整数")
        limits = {
            "llm_judge.global_max_concurrency": (1, 64, "全局 LLM 并发须为 1 至 64 的整数"),
            "llm_judge.per_task_max_concurrency": (1, 64, "单任务 LLM 并发须为 1 至 64 的整数"),
            "llm_judge.request_timeout_seconds": (10, 900, "LLM 请求超时须为 10 至 900 秒的整数"),
            "llm_judge.max_retries": (0, 2, "LLM 重试次数须为 0 至 2 的整数"),
        }
        values: dict[str, int] = {}
        for key, (minimum, maximum, message) in limits.items():
            value = payload.get(key)
            if value is None:
                continue
            try:
                number = int(value)
            except (TypeError, ValueError):
                raise ValueError(message)
            if str(value).strip() != str(number) or not minimum <= number <= maximum:
                raise ValueError(message)
            values[key] = number
        global_limit = values.get("llm_judge.global_max_concurrency", 30)
        if values.get("llm_judge.per_task_max_concurrency", 0) > global_limit:
            raise ValueError("单任务 LLM 并发不能超过全局 LLM 并发")
        medium = payload.get("testcase_verification.self_history_medium_rate")
        high = payload.get("testcase_verification.self_history_high_rate")
        if medium is None and high is None:
            return
        medium_v = float(medium if medium is not None else 0.20)
        high_v = float(high if high is not None else 0.50)
        if not (0 <= medium_v <= high_v <= 1):
            raise ValueError("测试用例查重阈值需满足 0 ≤ 中风险 ≤ 高风险 ≤ 1")

    def _parse_testcase_upload(self, file_item) -> list[dict]:
        from src.core.testcase_excel import load_daily_sheet_rows, rows_to_items
        name = str(getattr(file_item, "filename", "") or "").lower()
        if not name.endswith(".xlsx"):
            raise ValueError("测试用例文件须为 .xlsx")
        rows = load_daily_sheet_rows(file_item.data)
        if not rows:
            raise ValueError("日报 sheet 中没有有效测试用例")
        return rows_to_items(rows)

    def _resolve_snapshot_group(self, user, fields: dict) -> tuple[str, str, str, str]:
        if user.role == ROLE_GROUP_LEADER:
            if not user.group_id:
                raise ValueError("组长未绑定小组，无法上传汇总表")
            return user.group_id, user.group_name, user.department_id, user.department_name
        group_id = str(fields.get("group_id") or "").strip()
        if not group_id:
            raise ValueError("请选择小组")
        groups = self.db.list_groups(user.department_id)
        match = next((g for g in groups if g.get("group_id") == group_id), None)
        if not match:
            raise ValueError("小组不存在或不在本部门范围内")
        return match["group_id"], match.get("name", ""), user.department_id, user.department_name

    def _report_from_payload(self, user, payload: dict, report_date: str | None = None, created_at: str | None = None) -> Report:
        missing = [f for f in REQUIRED_REPORT_FIELDS if not str(payload.get(f, "")).strip()]
        if missing: raise ValueError(f"缺少必填字段: {', '.join(missing)}")
        if str(payload.get("submitted_at") or "").strip():
            raise ValueError("提交时间由服务器按上海时间自动生成，不能手动指定")
        derived_manager = self.db.get_derived_manager(user)
        manager = derived_manager.display_name if derived_manager else str(payload.get("manager") or "").strip()
        if getattr(user, "is_test_account", False) and not manager:
            manager = "测试负责人"
        if user.role != ROLE_MINISTER and not manager:
            raise ValueError("直属领导尚未建立，请手工填写")
        effective_report_date = report_date or self._new_report_date()
        actual_created_at = created_at or now_iso()
        report = Report("", user.user_id, user.display_name, user.department_id, user.department_name,
                        effective_report_date, manager, str(payload.get("title_summary") or ""),
                        str(payload.get("work_description") or ""), code_items=payload.get("code_items") or [],
                        document_items=payload.get("document_items") or [], testcase_items=payload.get("testcase_items") or [],
                        created_at=actual_created_at, group_id=user.group_id, group_name=user.group_name,
                        submitter_role=user.role)
        from src.core.pipeline import DailyReportPipeline
        report.report_id = DailyReportPipeline.build_report_id(report)
        body = build_report_docx(_report_to_payload(report))
        stored = self.storage.save_bytes("submitted_reports", f"SR_{report.report_date}_{report.employee_name}.docx", body, report.report_date)
        report.original_filename, report.stored_path, report.file_size, report.content_hash = stored.original_filename, stored.stored_path, stored.file_size, stored.content_hash
        return report

    def _submission_rules(self) -> dict:
        get_settings = getattr(self.db, "get_settings", None)
        if not callable(get_settings):
            return self.config.get("daily_report_submission", {})
        effective = apply_flat_settings(self.config, effective_settings(self.config, self.db))
        return effective.get("daily_report_submission", {})

    def _submission_now(self) -> datetime:
        return datetime.now(SHANGHAI_TIMEZONE)

    def _rollover_time(self) -> tuple[int, int]:
        rules = self._submission_rules()
        cutoff = str(rules.get("rollover_time", rules.get("cutoff_time", "09:00")))
        hour, minute = (int(part) for part in cutoff.split(":", 1))
        return hour, minute

    def _new_report_date(self, now: datetime | None = None) -> str:
        current = now or self._submission_now()
        if not self._submission_rules().get("enabled", True):
            return current.date().isoformat()
        hour, minute = self._rollover_time()
        if (current.hour, current.minute, current.second) < (hour, minute, 0):
            return (current.date() - timedelta(days=1)).isoformat()
        return current.date().isoformat()

    def _ensure_report_editable(self, report: Report, now: datetime | None = None) -> None:
        current = now or self._submission_now()
        if not self._submission_rules().get("enabled", True):
            return
        today = current.date().isoformat()
        previous = (current.date() - timedelta(days=1)).isoformat()
        hour, minute = self._rollover_time()
        if report.report_date == today:
            return
        if report.report_date == previous and (current.hour, current.minute, current.second) < (hour, minute, 0):
            return
        raise ValueError("该日报已超过可编辑时间")

    def _save_artifact(self, report: Report, user, file_item, artifact_type: str) -> str:
        allowed_key = "allowed_code_ext" if artifact_type == "code" else "allowed_document_ext"
        configured = self.config.get("upload", {}).get(allowed_key, [])
        allowed = {str(ext).lower() if str(ext).startswith(".") else f".{str(ext).lower()}" for ext in configured}
        suffix = Path(str(file_item.filename or "")).suffix.lower()
        if allowed and suffix not in allowed:
            raise ValueError(f"不支持的{artifact_type}附件类型：{suffix or '无扩展名'}")
        self._validate_upload_size(len(file_item.data))
        stored = self.storage.save_bytes(f"artifacts/{artifact_type}", file_item.filename, file_item.data, report.report_date)
        return self.db.save_artifact({"report_id": report.report_id, "owner_user_id": user.user_id, "department_id": user.department_id, "artifact_type": artifact_type, "original_filename": file_item.filename, "stored_path": stored.stored_path, "content_type": file_item.content_type, "file_size": stored.file_size, "content_hash": stored.content_hash})

    def _read_json_body(self) -> dict:
        length = int(self.headers.get("Content-Length", "0")); self._validate_upload_size(length)
        raw = self.rfile.read(length).decode("utf-8")
        return json.loads(raw or "{}")

    def _read_multipart_or_json(self):
        ctype = self.headers.get("Content-Type", "")
        if not ctype.startswith("multipart/form-data"):
            return self._read_json_body(), {}
        length = int(self.headers.get("Content-Length", "0")); self._validate_upload_size(length)
        body = self.rfile.read(length)
        msg = BytesParser(policy=default).parsebytes(f"Content-Type: {ctype}\r\n\r\n".encode() + body)
        fields, files = {}, {}
        for part in msg.iter_parts():
            name = part.get_param("name", header="content-disposition")
            filename = part.get_filename()
            data = part.get_payload(decode=True) or b""
            if filename:
                files.setdefault(name, []).append(type("Upload", (), {"filename": filename, "data": data, "content_type": part.get_content_type()}))
            else:
                fields[name] = data.decode(part.get_content_charset() or "utf-8")
        return fields, files

    def _validate_upload_size(self, length: int) -> None:
        limit = max(1, int(self.config.get("upload", {}).get("max_file_mb", 50))) * 1024 * 1024
        if int(length) > limit:
            raise RequestTooLarge(f"请求或单个文件超过 {limit // 1024 // 1024} MB 限制")

    def _send_json(self, data, status: HTTPStatus = HTTPStatus.OK, headers: dict[str, str] | None = None) -> None:
        body = json.dumps(data, ensure_ascii=False).encode("utf-8")
        self._finish_audit(int(status), data)
        self.send_response(status); self.send_header("Content-Type", "application/json; charset=utf-8"); self.send_header("Content-Length", str(len(body))); self.send_header("Cache-Control", "no-store")
        for k, v in (headers or {}).items(): self.send_header(k, v)
        self.end_headers(); self.wfile.write(body)

    def _send_file(self, path: Path, download_name: str | None = None) -> None:
        self._send_bytes(path.read_bytes(), mimetypes.guess_type(path.name)[0] or "application/octet-stream", download_name)

    def _send_bytes(self, body: bytes, content_type: str, download_name: str | None = None) -> None:
        self._finish_audit(int(HTTPStatus.OK))
        self.send_response(HTTPStatus.OK); self.send_header("Content-Type", content_type); self.send_header("Content-Length", str(len(body))); self.send_header("Cache-Control", "no-store")
        if download_name: self.send_header("Content-Disposition", attachment_header(download_name))
        self.end_headers(); self.wfile.write(body)

    def _report_summary(self, report: Report, viewer=None) -> dict:
        return {"report_id": report.report_id, "owner_user_id": report.owner_user_id,
                "employee_name": report.employee_name, "department_name": report.department_name,
                "group_name": report.group_name, "submitter_role": report.submitter_role,
                "report_date": report.report_date, "title_summary": report.title_summary,
                "original_filename": report.original_filename, "created_at": report.created_at, "updated_at": report.updated_at,
                "status": report.status, "can_edit": bool(viewer and viewer.user_id == report.owner_user_id and self._is_report_editable(report)),
                "code_item_count": len(report.code_items), "document_item_count": len(report.document_items),
                "artifacts": self.db.list_artifacts_for_report(report.report_id)}

    def _report_detail(self, report: Report) -> dict:
        return report.to_dict() | {"artifacts": self.db.list_artifacts_for_report(report.report_id)}

    def _is_report_editable(self, report: Report) -> bool:
        try:
            self._ensure_report_editable(report)
            return True
        except ValueError:
            return False

    def _clear_business_storage(self) -> int:
        root = Path(self.config["storage"]["root_dir"]).resolve()
        removed = 0
        for category in ("submitted_reports", "artifacts", "testcase_snapshots", "tasks", "exports", "backups"):
            path = (root / category).resolve()
            if root not in path.parents or not path.exists():
                continue
            removed += sum(1 for item in path.rglob("*") if item.is_file())
            shutil.rmtree(path)
        raw_log = Path(str(self.config.get("llm_judge", {}).get("raw_log_path") or "")).resolve()
        if raw_log.exists() and root in raw_log.parents:
            raw_log.unlink(); removed += 1
        return removed

    def _storage_stats(self) -> dict:
        root = Path(self.config["storage"]["root_dir"])
        size = sum(p.stat().st_size for p in root.rglob("*") if p.is_file()) if root.exists() else 0
        return {"root_dir": str(root), "bytes": size, "mb": round(size / 1024 / 1024, 2)}

def _report_to_payload(report: Report) -> dict:
    return {"name": report.employee_name, "department": report.department_name, "report_date": report.report_date, "submitted_at": report.created_at, "manager": report.manager, "title_summary": report.title_summary, "work_description": report.work_description, "code_items": report.code_items, "document_items": report.document_items, "testcase_items": report.testcase_items}


def _build_reports_docx(reports: list[Report]) -> bytes:
    from io import BytesIO
    from docx import Document
    doc = Document(); doc.add_heading("周报汇总", level=1); doc.add_paragraph(f"共 {len(reports)} 份日报")
    table = doc.add_table(rows=1, cols=8); table.style = "Table Grid"
    for i, h in enumerate(["查重归属日期", "姓名", "职务", "部门", "小组", "标题", "状态", "提交时间"]): table.rows[0].cells[i].text = h
    for r in reports:
        row = table.add_row().cells
        for i, v in enumerate([r.report_date, r.employee_name, r.submitter_role, r.department_name, r.group_name, r.title_summary, r.status, r.created_at]): row[i].text = str(v or "")
    buf = BytesIO(); doc.save(buf); return buf.getvalue()


def _periodic_export_args(query: dict[str, list[str]]) -> tuple[str, str, str]:
    report_type = str((query.get("report_type") or [""])[0]).strip()
    start = str((query.get("start") or [""])[0]).strip()
    end = str((query.get("end") or [""])[0]).strip()
    if report_type not in {"weekly", "monthly"}:
        raise ValueError("报表类型必须是 weekly 或 monthly")
    if not start or not end:
        raise ValueError("请选择开始日期和结束日期")
    start, end = parse_date(start), parse_date(end)
    if start > end:
        raise ValueError("开始日期不能晚于结束日期")
    return report_type, start, end


def _export_scope_label(user) -> str:
    return {"employee": "本人", "group_leader": "本组", "director": "本部门", "minister": "全公司"}.get(user.role, "权限范围")


def _periodic_missing_reports(db, user, start: str, end: str, reports: list[Report]) -> list[dict]:
    submitted = {(report.owner_user_id, report.report_date) for report in reports}
    submitters = db.list_export_submitters(user)
    current, final = datetime.strptime(start, "%Y-%m-%d").date(), datetime.strptime(end, "%Y-%m-%d").date()
    workdays = []
    while current <= final:
        if current.weekday() < 5:
            workdays.append(current.isoformat())
        current += timedelta(days=1)
    missing = []
    for day in workdays:
        for submitter in submitters:
            if (submitter["user_id"], day) not in submitted:
                missing.append({"日期": day, "姓名": submitter.get("display_name", ""), "职务": submitter.get("role", ""), "部门": submitter.get("department_name", ""), "小组": submitter.get("group_name", ""), "日报": "无"})
    return missing


def _display_value(value) -> str:
    if value is None:
        return ""
    if isinstance(value, (dict, list)):
        return json.dumps(value, ensure_ascii=False)
    return str(value)


def _detail_headers(rows: list[dict]) -> list[str]:
    headers: list[str] = []
    for row in rows:
        for key in row:
            if key not in headers:
                headers.append(key)
    return headers


def _add_detail_table(doc, title: str, rows: list[dict]) -> None:
    doc.add_heading(title, level=3)
    if not rows:
        doc.add_paragraph("无")
        return
    headers = _detail_headers(rows)
    table = doc.add_table(rows=1, cols=len(headers)); table.style = "Table Grid"
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = _display_value(header)
    for item in rows:
        cells = table.add_row().cells
        for index, header in enumerate(headers):
            cells[index].text = _display_value(item.get(header, ""))


def _build_periodic_reports_docx(title: str, scope: str, start: str, end: str, reports: list[Report], artifacts: dict[str, list[dict]], missing_reports: list[dict] | None = None) -> bytes:
    from io import BytesIO
    from docx import Document

    doc = Document()
    doc.add_heading(title, level=1)
    missing_reports = missing_reports or []
    doc.add_paragraph(f"范围：{scope}；日期：{start} 至 {end}；共 {len(reports)} 份日报，工作日无日报 {len(missing_reports)} 条")
    if not reports and not missing_reports:
        doc.add_paragraph("暂无日报")
    for report in sorted(reports, key=lambda item: (item.report_date, item.employee_name, item.created_at)):
        doc.add_heading(f"查重归属：{report.report_date} · {report.employee_name}", level=2)
        doc.add_paragraph(f"职务：{report.submitter_role}；部门/小组：{report.department_name or '—'} / {report.group_name or '—'}")
        doc.add_paragraph(f"提交时间：{report.created_at or '—'}；最后修改：{report.updated_at or '—'}")
        doc.add_paragraph(f"标题：{report.title_summary}")
        doc.add_paragraph(f"工作描述：{report.work_description}")
        _add_detail_table(doc, "代码记录", report.code_items)
        _add_detail_table(doc, "文档记录", report.document_items)
        _add_detail_table(doc, "测试用例记录", report.testcase_items)
        _add_detail_table(doc, "附件", [{"文件名": item.get("original_filename", ""), "类型": item.get("artifact_type", ""), "提交时间": item.get("created_at", "")} for item in artifacts.get(report.report_id, [])])
    _add_detail_table(doc, "工作日无日报", missing_reports)
    buf = BytesIO(); doc.save(buf); return buf.getvalue()


def _build_periodic_reports_xlsx(title: str, scope: str, start: str, end: str, reports: list[Report], artifacts: dict[str, list[dict]], missing_reports: list[dict] | None = None) -> bytes:
    from io import BytesIO
    from openpyxl import Workbook

    wb = Workbook()
    summary = wb.active; summary.title = "日报明细"
    summary.append(["报表", title, "范围", scope, "开始日期", start, "结束日期", end])
    summary.append(["日报ID", "查重归属日期", "姓名", "职务", "部门", "小组", "标题", "工作描述", "提交时间", "最后修改时间", "状态"])
    ordered = sorted(reports, key=lambda item: (item.report_date, item.employee_name, item.created_at))
    for report in ordered:
        summary.append([report.report_id, report.report_date, report.employee_name, report.submitter_role, report.department_name, report.group_name, report.title_summary, report.work_description, report.created_at, report.updated_at, "已提交"])
    for missing in missing_reports or []:
        summary.append(["", missing["日期"], missing["姓名"], missing["职务"], missing["部门"], missing["小组"], "无", "无", "", "", "无"])

    def append_detail_sheet(name: str, rows: list[dict]) -> None:
        sheet = wb.create_sheet(name)
        headers = _detail_headers(rows)
        sheet.append(headers or ["暂无记录"])
        for row in rows:
            sheet.append([_display_value(row.get(header, "")) for header in headers])

    code_rows, document_rows, testcase_rows, artifact_rows = [], [], [], []
    for report in ordered:
        prefix = {"日报ID": report.report_id, "查重归属日期": report.report_date, "提交时间": report.created_at, "姓名": report.employee_name}
        code_rows.extend([prefix | item for item in report.code_items])
        document_rows.extend([prefix | item for item in report.document_items])
        testcase_rows.extend([prefix | item for item in report.testcase_items])
        artifact_rows.extend([prefix | {"文件名": item.get("original_filename", ""), "类型": item.get("artifact_type", ""), "提交时间": item.get("created_at", "")} for item in artifacts.get(report.report_id, [])])
    append_detail_sheet("代码记录", code_rows)
    append_detail_sheet("文档记录", document_rows)
    append_detail_sheet("测试用例", testcase_rows)
    append_detail_sheet("附件", artifact_rows)
    buf = BytesIO(); wb.save(buf); return buf.getvalue()


def run_server(config_path: str, host: str | None = None, port: int | None = None) -> None:
    config = load_config(config_path); storage = FileStorage(config); db = create_database(config)
    DailyReportHandler.config = config; DailyReportHandler.storage = storage; DailyReportHandler.db = db
    bind_host = host or str(config.get("app", {}).get("host", "0.0.0.0")); bind_port = int(port or config.get("app", {}).get("port", 8000))
    server = ThreadingHTTPServer((bind_host, bind_port), DailyReportHandler)
    runtime_log(db, "web", "started", f"Daily Report V2 backend running at http://{bind_host}:{bind_port}", context={"host": bind_host, "port": bind_port})
    try:
        server.serve_forever()
    finally:
        runtime_log(db, "web", "stopped", "Daily Report V2 backend 已停止")
        server.server_close()


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--config", default=None, help="Path to config.yaml; defaults to project config/config.yaml")
    parser.add_argument("--host", default=None)
    parser.add_argument("--port", type=int, default=None)
    args = parser.parse_args()
    config_path = args.config or str(default_config_path())
    run_server(config_path, args.host, args.port)


if __name__ == "__main__": main()
