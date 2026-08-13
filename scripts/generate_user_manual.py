"""Generate the daily_report user manual as a Word document."""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT


def set_cell_shading(cell, color: str) -> None:
    shading = cell._element.get_or_add_tcPr()
    shd = shading.makeelement(qn("w:shd"), {
        qn("w:fill"): color,
        qn("w:val"): "clear",
    })
    shading.append(shd)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], header_fill: str = "D9E2F3") -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        set_cell_shading(cell, header_fill)
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    doc.add_paragraph()


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_para(doc: Document, text: str, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(11)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        for r in p.runs:
            r.font.size = Pt(11)


def build_document() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.8)

    # Cover
    for _ in range(6):
        doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run("日报核验系统\n用户使用与技术说明文档")
    tr.bold = True
    tr.font.size = Pt(22)
    tr.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run("daily_report · 企业内部日报采集与智能核验平台")
    sr.font.size = Pt(14)
    sr.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mr = meta.add_run("文档版本：V2.0\n编写日期：2026年7月10日\n读者对象：员工、组长、主任、部长、系统管理员及运维人员")
    mr.font.size = Pt(11)

    doc.add_page_break()

    # TOC
    add_heading(doc, "目  录", 1)
    toc_items = [
        "第一章  项目背景与建设目标",
        "第二章  系统概述",
        "第三章  角色与权限说明",
        "第四章  核心功能详解",
        "第五章  分角色使用指南",
        "第六章  技术实现路线",
        "第七章  系统访问与部署说明",
        "第八章  常见问题与术语表",
        "附录 A  默认测试账号",
        "附录 B  风险等级与处理建议",
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.left_indent = Cm(0.5)
        for r in p.runs:
            r.font.size = Pt(12)
    doc.add_page_break()

    # Chapter 1
    add_heading(doc, "第一章  项目背景与建设目标", 1)
    add_heading(doc, "1.1  项目背景", 2)
    add_para(doc, "在企业日常管理中，员工需按日提交工作日报，用于记录工作内容、测试进展、代码与文档产出。传统方式存在以下问题：")
    add_bullets(doc, [
        "日报质量参差不齐：部分日报内容过短、信息不足，难以反映真实工作。",
        "重复填报风险：员工可能复制粘贴历史日报，或与他人内容高度相似。",
        "测试用例与日报脱节：员工在日报中声称「新增/修改测试用例」，但难以与部门实际测试总表核对。",
        "附件管理分散：代码、文档等附件缺乏统一归档与核验机制。",
        "领导复核效率低：人工逐份阅读日报，难以快速识别高风险条目。",
    ])
    add_para(doc, "为解决上述问题，本项目建设了「日报核验系统（daily_report）」，实现日报在线录入、自动生成 Word 文档、多维度自动核验、查重任务管理与结果导出，支撑部门级与全公司级的日报质量管控。")

    add_heading(doc, "1.2  建设目标", 2)
    add_table(doc, ["目标", "说明"], [
        ["规范日报提交", "统一表单结构，自动生成标准 Word 日报"],
        ["自动质量检测", "识别内容缺失、过短、信息量偏低等问题"],
        ["智能查重核验", "结合文本相似度、语义分析与可选大模型判断，发现重复或疑似抄袭"],
        ["测试用例对齐", "将员工日报中的测试工作与部门每日总表及历史基线比对"],
        ["分级权限管理", "员工、组长、主任、部长按组织树查看本人及下级数据"],
        ["可部署可运维", "支持本地运行、Docker 部署、内网大模型接入及数据备份"],
    ])

    # Chapter 2
    add_heading(doc, "第二章  系统概述", 1)
    add_heading(doc, "2.1  系统定位", 2)
    add_para(doc, "日报核验系统是一套企业内部日报采集与智能核验平台。业务人员通过浏览器填写日报；系统按小组和部门自动识别直属领导并生成 Word 文件；组长、主任或部长可在各自组织范围内发起查重任务并导出结果。")

    add_heading(doc, "2.2  主要功能模块", 2)
    add_table(doc, ["模块", "功能简述", "主要使用角色"], [
        ["日报录入", "填写日报、上传附件、自动带出直属领导", "员工、组长、主任、部长"],
        ["日报记录", "按姓名、日期、职务查看和导出可见日报", "员工、组长、主任、部长"],
        ["查重任务/记录", "创建私有任务、查看结果、导出报告", "组长、主任、部长"],
        ["测试总表", "上传或查看部门每日测试用例 Excel 快照", "组长、主任、部长"],
        ["系统管理", "账号、部门、小组、参数、运维和清理", "管理员"],
    ])

    add_heading(doc, "2.3  系统界面结构", 2)
    add_para(doc, "登录后，左侧导航栏根据角色显示不同菜单：")
    add_bullets(doc, [
        "日报录入、日报记录：员工、组长、主任、部长可见",
        "查重任务、查重记录：组长、主任、部长可见，任务记录仅创建者本人可见",
        "测试总表：组长、主任、部长可见（上传权限为组长和主任）",
        "系统管理：仅管理员可见",
    ])

    # Chapter 3
    add_heading(doc, "第三章  角色与权限说明", 1)
    add_para(doc, "业务组织按员工 → 组长 → 主任 → 部长逐级管理；管理员是独立的系统运维角色，不查看业务日报或查重结果。")

    add_heading(doc, "3.1  角色对照表", 2)
    add_table(doc, ["角色标识", "中文名称", "数据可见范围", "核心权限"], [
        ["employee", "员工", "仅本人日报", "录入、查看本人记录、撤回、导出本人周报"],
        ["group_leader", "组长", "本人及本组员工日报", "本组范围查重、上传本部门测试总表"],
        ["director", "主任", "本人及本部门全部日报", "本部门查重、上传本部门测试总表"],
        ["minister", "部长", "全公司日报", "全公司查重、查看各部门总表"],
        ["admin", "管理员", "不查看业务数据", "账号、部门、小组、参数、备份及业务数据清理"],
    ])

    add_heading(doc, "3.2  权限矩阵", 2)
    add_table(doc, ["功能", "员工", "组长", "主任", "部长", "管理员"], [
        ["提交/查看日报", "本人", "本人及本组", "本人及本部门", "全公司", "—"],
        ["撤回本人未锁定日报", "✓", "✓", "✓", "✓", "—"],
        ["创建查重任务", "—", "本组", "本部门", "全公司", "—"],
        ["查看/下载本人查重任务", "—", "✓", "✓", "✓", "—"],
        ["上传测试用例总表", "—", "本部门", "本部门", "—", "—"],
        ["账号/部门/小组/参数管理", "—", "—", "—", "—", "✓"],
        ["测试模型/备份/业务清理", "—", "—", "—", "—", "✓"],
    ])

    add_heading(doc, "3.3  重要权限规则", 2)
    add_bullets(doc, [
        "姓名、职务、部门和小组由登录身份自动填充且不可修改。",
        "直属领导自动推导：员工对应组长、组长对应主任、主任对应部长；领导尚未创建时才允许手工填写。",
        "日报撤回：仅当日报未被查重任务锁定且状态非「已撤回」时可撤回。",
        "查重后锁定：某份日报被纳入查重任务后会被锁定，不可再撤回或修改。",
        "测试总表上传：组长和主任只能上传本部门的总表。",
        "自注册：新用户选择部门和小组，注册后角色固定为「员工」。",
    ])

    # Chapter 4
    add_heading(doc, "第四章  核心功能详解", 1)

    add_heading(doc, "4.1  日报录入与 Word 生成", 2)
    add_heading(doc, "4.1.1  录入内容", 3)
    add_table(doc, ["字段", "是否必填", "说明"], [
        ["姓名", "自动填充", "员工账号自动带出，不可改"],
        ["部门", "自动填充", "员工账号自动带出，不可改"],
        ["日期", "必填", "日报所属工作日"],
        ["直属领导", "自动/必填", "按组织自动带出；组织未完善时手工填写"],
        ["标题总结", "必填", "当日工作一句话概括"],
        ["工作描述", "必填", "详细工作内容"],
        ["测试用例相关工作", "选填", "场景、类型、负责人、用例编号、说明"],
        ["代码相关工作", "选填", "文件名、语言、类型、说明；可上传代码附件"],
        ["文档相关工作", "选填", "文档名、类型、工作类型、说明；可上传文档附件"],
    ])

    add_heading(doc, "4.1.2  提交流程", 3)
    add_bullets(doc, [
        "填写表单，点击「提交为 Word 日报」。",
        "系统校验必填项。",
        "自动生成标准 Word 文档（.docx）。",
        "文件存入 storage/submitted_reports/，元数据写入数据库。",
        "代码/文档附件单独存入 storage/artifacts/。",
    ])

    add_heading(doc, "4.1.3  提交记录管理", 3)
    add_para(doc, "在「提交记录」页面可查看日期、标题、状态、附件、提交时间；下载 Word 日报及附件；撤回未被锁定的日报；按日期范围导出「周报 Word」汇总。")

    add_heading(doc, "4.2  查重任务与多维核验", 2)
    add_para(doc, "查重任务是系统的核心能力。领导选定日期范围、部门/员工范围后，系统对范围内每份日报执行多项检查。")

    add_heading(doc, "4.2.1  检查项一览", 3)
    add_table(doc, ["检查项", "中文名称", "检查内容"], [
        ["quality_check", "日报质量", "标题/描述缺失、描述过短、信息量偏低"],
        ["report_duplicate_check", "日报查重", "与本人历史及他人近期日报的相似度"],
        ["testcase_check", "测试用例核验", "日报中测试工作与部门总表及历史基线比对"],
        ["code_check", "代码记录", "记录代码条目与附件；支持 JPlag 代码查重"],
        ["document_check", "文档记录", "记录文档条目与附件；支持本人历史文档比对"],
    ])

    add_heading(doc, "4.2.2  日报查重算法", 3)
    add_para(doc, "系统采用多阶段候选召回 + 加权评分 + 可选大模型裁决的流程：")
    add_para(doc, "第一阶段：候选召回", bold=True)
    add_bullets(doc, [
        "本人历史范围：默认回溯 30 天内的本人历史日报。",
        "跨人范围：默认回溯 3 天内、同部门其他员工的日报。",
        "排除：已撤回日报、日期晚于当前日报的记录。",
    ])
    add_para(doc, "第二阶段：相似度评分（默认权重可在管理员页面调整）", bold=True)
    add_table(doc, ["维度", "默认权重", "计算方法"], [
        ["文本相似度", "45%", "字符级 SequenceMatcher"],
        ["语义相似度", "45%", "TF-IDF 字符 n-gram + 余弦相似度"],
        ["时间衰减", "10%", "距今天数越近，分数越高"],
    ])
    add_para(doc, "第三阶段：大模型裁决（可选）", bold=True)
    add_bullets(doc, [
        "取得分最高的 Top-N 条候选（默认 N=3）。",
        "相似度超过阈值（默认 0.72，低信息日报 0.82）时，调用内网 OpenAI 兼容大模型做最终重复判定。",
        "大模型可区分本人合理复述、跨人协作、疑似抄袭等情形。",
    ])

    add_heading(doc, "4.2.3  测试用例核验逻辑", 3)
    add_para(doc, "员工不上传测试 Excel。核验依赖组长或主任上传的部门每日总表：")
    add_bullets(doc, [
        "系统读取员工所属部门、日报日期对应的当天总表。",
        "读取该日期之前最近一份历史基线总表。",
        "将员工日报中填写的测试用例行与总表逐条比对。",
    ])
    add_table(doc, ["员工填报类型", "核验规则", "典型风险"], [
        ["新增", "当天总表应存在，历史基线不应存在", "声称新增但历史已有 → 中风险"],
        ["修改", "当天与历史基线应有差异", "声称修改但无变化 → 高风险"],
        ["任意", "当天总表应能找到对应用例", "总表中未找到 → 高风险"],
        ["—", "部门当天未上传总表", "未知风险"],
    ])

    add_heading(doc, "4.2.4  任务输出", 3)
    add_para(doc, "每次查重任务在 storage/tasks/<task_id>/ 独立保存：review_cases.json（完整结果）、review_cases.xlsx（Excel 汇总）、review_cases.docx（Word 报告）、run_metadata.json（运行配置）。任务完成后，涉及日报将被锁定。")

    add_heading(doc, "4.3  测试用例总表管理", 2)
    add_bullets(doc, [
        "上传者：组长或主任，且只能上传本人所在部门。",
        "文件格式：.xlsx。",
        "上传频率：建议每个工作日上传一份。",
        "命名规则：按「快照日期」归档，同一天可覆盖替换。",
        "总表 Excel 建议包含「用例编号」列，以便精确匹配。",
    ])

    add_heading(doc, "4.4  系统管理（管理员专属）", 2)
    add_para(doc, "管理员可进行账号、部门、小组、参数配置及运维操作，并通过两次确认清除业务数据。管理员不查看业务日报或查重结果。")
    add_table(doc, ["参数", "含义", "默认值"], [
        ["单任务并行数", "兼容参数；生产建议保持 1，避免 CPU 密集查重拖慢系统", "1"],
        ["日报进入大模型候选数 Top-N", "送入大模型裁决的候选条数", "3"],
        ["日报候选最低分", "进入大模型判断的相似度门槛", "0.72"],
        ["本人历史天数", "查重时回溯本人历史范围", "30 天"],
        ["跨人近几天", "查重时回溯他人日报范围", "3 天"],
        ["跨人范围", "department（同部门）或 all（全公司）", "department"],
        ["文本/语义/时间权重", "相似度加权系数", "0.45 / 0.45 / 0.10"],
        ["启用大模型", "是否调用内网 LLM", "false"],
        ["上传大小 MB", "单文件上传限制", "50"],
    ])

    # Chapter 5
    add_heading(doc, "第五章  分角色使用指南", 1)

    add_heading(doc, "5.1  员工使用指南", 2)
    add_heading(doc, "5.1.1  首次使用", 3)
    add_bullets(doc, [
        "打开系统地址（如 http://服务器IP:8000）。",
        "点击「注册员工账号」，填写账号、密码、姓名并选择部门和小组；或联系管理员开通账号。",
    ])
    add_heading(doc, "5.1.2  每日工作流程", 3)
    add_para(doc, "登录 → 员工录入 → 填写当日工作 → 提交 → 提交记录中确认")
    add_para(doc, "操作步骤：")
    add_bullets(doc, [
        "登录后默认进入「员工录入」页面。",
        "确认姓名、部门已自动填充。",
        "确认直属领导已按组织自动带出；如对应领导尚未创建，再手工填写。",
        "（可选）在测试用例、代码、文档相关工作中填写条目并上传附件。",
        "点击「提交为 Word 日报」，等待提示「已生成 xxx.docx」。",
        "切换到「提交记录」确认提交成功。",
    ])
    add_heading(doc, "5.1.3  注意事项", 3)
    add_bullets(doc, [
        "测试相关工作无需自行上传 Excel，系统会读取组长或主任上传的部门总表进行核验。",
        "提交后若发现错误，可在「提交记录」点击「撤回」（未被查重锁定时）。",
        "需要周报时，在「提交记录」选择起止日期，点击「导出周报 Word」。",
    ])

    add_heading(doc, "5.2  组长与主任使用指南", 2)
    add_heading(doc, "5.2.1  每日必做：上传测试总表", 3)
    add_bullets(doc, [
        "登录后进入「测试总表」。",
        "选择「快照日期」（一般为当日）。",
        "选择本部门测试用例 Excel 总表文件。",
        "点击「上传/替换当天总表」，在列表中确认上传成功。",
        "说明：若当天未上传总表，员工日报中的测试用例核验将标记为「未知风险」。",
    ])
    add_heading(doc, "5.2.2  定期查重复核", 3)
    add_bullets(doc, [
        "进入「查重任务」，填写开始/结束日期。",
        "「部门范围」自动锁定为本部门。",
        "「员工范围」可留空（查全部）或填写特定员工姓名。",
        "点击「创建查重任务」，等待任务完成。",
        "在任务列表点击「查看」浏览结果，或使用「Excel」「Word」下载报告。",
        "使用搜索框、日期筛选、风险等级筛选定位高风险条目。",
        "「快速运行」可对默认范围立即执行一次查重，适合日常快速抽检。",
    ])
    add_heading(doc, "5.2.3  结果阅读方法", 3)
    add_para(doc, "每条结果卡片包含员工姓名、部门、日期、标题、总体风险徽章及各检查项详情。建议处理优先级：高风险 → 中风险 → 未知（缺数据）→ 低风险。")

    add_heading(doc, "5.3  部长使用指南", 2)
    add_table(doc, ["能力", "组长", "主任", "部长"], [
        ["日报可见范围", "本人及本组", "本人及本部门", "全公司"],
        ["查重执行范围", "本人及本组", "本部门", "全公司或指定部门"],
        ["上传测试总表", "✓", "✓", "✗（可查看）"],
        ["任务记录", "仅本人创建", "仅本人创建", "仅本人创建"],
    ])
    add_para(doc, "全公司查重操作：进入「查重任务」→ 设置日期范围 → 「部门范围」选择「全公司」或指定部门 → 创建任务并审阅结果 → 导出 Excel/Word 报告。")

    add_heading(doc, "5.4  管理员使用指南", 2)
    add_heading(doc, "5.4.1  系统初始化清单", 3)
    add_bullets(doc, [
        "修改所有默认测试账号密码。",
        "创建正式部门和其下属小组。",
        "为员工、组长、主任和部长建立账号及组织归属；直属领导由系统推导。",
        "确保每个小组仅一名组长、每部门仅一名主任、全公司仅一名部长。",
        "配置内网大模型地址并测试连接。",
        "根据企业要求调整查重阈值与 Top-N 参数。",
        "确认 JPlag jar 已部署（如需代码查重）。",
        "建立定期 storage 备份计划。",
    ])
    add_heading(doc, "5.4.2  大模型配置步骤", 3)
    add_bullets(doc, [
        "在「参数配置」中设置：启用大模型 = true，模型 Base URL，模型名。",
        "将 API Key 写入 /opt/daily_report/config/llm_api_key，并设置权限 chmod 600。",
        "点击「测试大模型连接」确认返回成功。",
    ])

    # Chapter 6
    add_heading(doc, "第六章  技术实现路线", 1)
    add_heading(doc, "6.1  总体架构", 2)
    add_para(doc, "浏览器前端（HTML/CSS/JS）通过 HTTP 与 Python 后端服务通信；后端连接 SQLite 元数据库、本地文件存储（storage/）及可选的内网 OpenAI 兼容大模型。生产环境推荐 Nginx 反向代理 + Docker 容器化部署。")

    add_heading(doc, "6.2  技术栈", 2)
    add_table(doc, ["层次", "技术选型", "说明"], [
        ["前端", "原生 HTML / CSS / JavaScript", "无框架，轻量易部署"],
        ["后端", "Python 3.12", "标准库 HTTP 服务 + 业务模块"],
        ["数据库", "SQLite", "存用户、日报元数据、任务记录；可扩展 PostgreSQL"],
        ["文件存储", "本地文件系统", "Word、Excel、附件、任务结果"],
        ["文档处理", "python-docx、openpyxl", "生成/读取 Word 与 Excel"],
        ["文本分析", "scikit-learn、difflib", "TF-IDF 语义相似度、序列匹配"],
        ["代码查重", "JPlag v6.0.0 + JDK 21", "可选，jar 缺失时降级为记录模式"],
        ["大模型", "OpenAI Compatible API", "内网部署，支持重复判定"],
    ])

    add_heading(doc, "6.3  目录结构", 2)
    add_para(doc, "daily_report/backend/（Python 后端）、frontend/（前端静态页面）、deploy/（Docker、Nginx、部署脚本）、vendor/jplag/（JPlag 查重引擎）、storage/（运行时数据目录）。")

    add_heading(doc, "6.4  核心处理流水线", 2)
    add_bullets(doc, [
        "范围过滤：按角色、日期、部门、员工筛选目标日报。",
        "质量检查：校验字段完整性与信息量。",
        "日报查重：候选召回 → 加权评分 → 大模型裁决。",
        "测试用例核验：读取部门快照与历史基线比对。",
        "代码/文档记录：归档附件并标记 JPlag/历史查重状态。",
        "汇总输出：计算总体风险，生成 JSON/Excel/Word。",
        "锁定日报：防止查重后数据被撤回。",
    ])

    add_heading(doc, "6.5  安全与权限设计", 2)
    add_bullets(doc, [
        "Session 认证：登录后 Cookie 保存 Session ID，默认 7 天有效。",
        "接口级鉴权：每个 API 按角色校验。",
        "组织隔离：员工仅见本人；组长见本组；主任见本部门；部长见全公司。",
        "任务隔离：查重任务、结果和下载仅创建者本人可访问。",
        "密码存储：哈希存储，不明文保存。",
    ])

    # Chapter 7
    add_heading(doc, "第七章  系统访问与部署说明", 1)
    add_table(doc, ["环境", "地址"], [
        ["本地开发", "http://127.0.0.1:8000"],
        ["生产部署", "由运维配置的域名或 IP（通常经 Nginx 80/443 端口）"],
    ])
    add_para(doc, "本地快速启动：进入 daily_report/backend，创建虚拟环境，pip install -r requirements.txt，执行 python -m src.web.app --config config/config.yaml --host 0.0.0.0 --port 8000。详细生产部署步骤见项目内 DEPLOY.md。")

    # Chapter 8
    add_heading(doc, "第八章  常见问题与术语表", 1)
    add_heading(doc, "8.1  常见问题（FAQ）", 2)
    faq = [
        ("提交日报时提示缺少必填字段？", "请确认已填写日期、直属负责人、标题总结、工作描述。"),
        ("为什么无法撤回日报？", "该日报可能已被查重任务锁定，或状态已为「已撤回」。"),
        ("测试用例核验显示「未知风险」？", "通常因为组长或主任尚未上传当天的部门测试总表，请联系对应负责人补传。"),
        ("查重结果中「大模型调用失败」？", "请联系管理员检查大模型地址、API Key 及网络连通性。"),
        ("代码查重未生效？", "确认 vendor/jplag/jplag.jar 已正确部署，且服务器已安装 JDK 21。"),
        ("忘记密码怎么办？", "联系系统管理员重置密码。"),
    ]
    for q, a in faq:
        add_para(doc, f"问：{q}", bold=True)
        add_para(doc, f"答：{a}")

    add_heading(doc, "8.2  术语表", 2)
    add_table(doc, ["术语", "解释"], [
        ["日报", "员工按日提交的工作记录 Word 文档"],
        ["查重任务", "对一批日报执行自动核验的后台作业"],
        ["快照/总表", "组长或主任上传的部门每日测试用例 Excel"],
        ["基线", "某日报日期之前最近一份部门测试总表"],
        ["Top-N", "相似度最高的前 N 条候选日报"],
        ["锁定", "日报被纳入查重任务后不可撤回的状态"],
    ])

    # Appendix
    add_heading(doc, "附录 A  默认测试账号（生产环境务必修改密码）", 1)
    add_table(doc, ["角色", "账号", "密码", "说明"], [
        ["管理员", "admin", "admin123", "系统管理"],
        ["部长", "minister", "minister123", "全公司日报与查重"],
        ["主任", "director", "director123", "研发部主任"],
        ["组长", "leader", "leader123", "研发一组组长"],
        ["员工", "employee", "employee123", "研发部员工张三"],
        ["员工", "employee2", "employee234", "研发部员工李四"],
    ])

    add_heading(doc, "附录 B  风险等级与处理建议", 1)
    add_table(doc, ["等级", "标识", "建议处理"], [
        ["高风险", "红色", "优先约谈核实，必要时要求补充说明或重新提交"],
        ["中风险", "橙色", "安排复核，关注测试用例类型不符或中度相似"],
        ["低风险", "黄色", "记录备案，抽查即可"],
        ["正常", "绿色", "无需特别处理"],
        ["未知", "灰色", "检查是否缺数据（总表未上传、大模型不可用等）"],
    ])

    return doc


def main() -> None:
    root = Path(__file__).resolve().parents[1]
    doc = build_document()
    outputs = [root / "daily_report_user_manual.docx"]
    for out in outputs:
        doc.save(str(out))
        print(f"Generated: {out}")


if __name__ == "__main__":
    main()
