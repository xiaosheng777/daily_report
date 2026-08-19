from datetime import datetime
from types import SimpleNamespace
from zoneinfo import ZoneInfo
from io import BytesIO

from docx import Document
from openpyxl import load_workbook
import pytest
from src.core.docx_io import build_report_docx
from src.core.text import low_information
from src.db.sqlite_backend import ROLE_MINISTER
from src.web.app import DailyReportHandler, _build_periodic_reports_docx, _build_periodic_reports_xlsx, _periodic_export_args, _periodic_missing_reports, attachment_header, audit_changes, csv_safe, effective_settings, flatten_settings, redact_audit_value


def test_attachment_header_encodes_non_ascii_filename():
    header = attachment_header("SR_2026-06-30_张三.docx")

    header.encode("latin-1")
    assert 'filename="SR_2026-06-30_.docx"' in header
    assert "filename*=UTF-8''SR_2026-06-30_%E5%BC%A0%E4%B8%89.docx" in header


def test_audit_redaction_csv_protection_and_route_scope():
    changes = audit_changes(
        {"role": "employee", "api_key": "old", "password": "old"},
        {"role": "admin", "api_key": "new", "password": "new"},
        ("role", "api_key", "password"),
    )
    assert changes["role"] == {"before": "employee", "after": "admin"}
    assert changes["api_key"] == {"before": "[redacted]", "after": "[redacted]"}
    assert changes["password"] == {"before": "[redacted]", "after": "[changed]"}
    assert redact_audit_value("authorization", "Bearer secret") == "[redacted]"
    assert csv_safe("=1+1") == "'=1+1"
    assert DailyReportHandler._audit_descriptor("GET", "/api/my/reports") is None
    assert DailyReportHandler._audit_descriptor("GET", "/api/my/reports/r1/download") is None
    assert DailyReportHandler._audit_descriptor("GET", "/api/tasks/t1/export.xlsx")["action"] == "export.task_result"


def test_failed_audited_request_is_warning_and_audit_failure_does_not_raise():
    handler = object.__new__(DailyReportHandler)
    recorded = []
    handler.db = SimpleNamespace(append_audit_log=lambda payload: recorded.append(payload))
    handler.headers = {"X-Forwarded-For": "10.0.0.1", "User-Agent": "pytest"}
    handler.client_address = ("127.0.0.1", 1)
    handler._pending_audit = {"action": "auth.login", "category": "authentication", "summary": "登录", "severity": "info"}
    handler._finish_audit(401, {"error": "invalid_credentials"})
    assert recorded[0]["outcome"] == "failure"
    assert recorded[0]["severity"] == "warning"
    assert recorded[0]["client_ip"] == "10.0.0.1"

    handler._pending_audit = {"action": "admin.llm.test", "category": "administration", "summary": "测试 LLM 连接", "severity": "info"}
    handler._finish_audit(200, {"ok": False, "error": "Connection refused"})
    assert recorded[1]["outcome"] == "failure"
    assert recorded[1]["severity"] == "warning"
    assert recorded[1]["summary"] == "测试 LLM 连接失败"
    assert recorded[1]["error_type"] == "Connection refused"

    handler.db = SimpleNamespace(append_audit_log=lambda payload: (_ for _ in ()).throw(RuntimeError("db unavailable")))
    handler._pending_audit = {"action": "task.create", "category": "business", "summary": "创建查重任务", "severity": "info"}
    handler._finish_audit(201, {})


def test_low_information_uses_only_combined_text_length():
    assert low_information("完成", "接口") is True
    assert low_information("完成", "接口开发联调测试") is False


def test_employee_report_exports_are_paused_but_manager_exports_remain_available():
    handler = object.__new__(DailyReportHandler)
    sent = []
    handler._send_json = lambda body, status: sent.append((body, status))
    employee = SimpleNamespace(role="employee")
    manager = SimpleNamespace(role="group_leader")
    handler._require_business_user = lambda: employee
    assert handler._require_report_exporter() is None
    assert sent[0][0]["error"] == "report_export_paused_for_employees"
    handler._require_business_user = lambda: manager
    assert handler._require_report_exporter() is manager


def test_report_date_is_derived_from_server_time_and_payload_date_is_ignored(monkeypatch):
    handler = object.__new__(DailyReportHandler)
    handler.config = {"daily_report_submission": {"enabled": True, "timezone": "Asia/Shanghai", "rollover_time": "09:00"}}
    handler.db = SimpleNamespace(get_derived_manager=lambda user: None)
    handler.storage = SimpleNamespace(save_bytes=lambda *args: SimpleNamespace(
        original_filename="report.docx", stored_path="submitted_reports/report.docx", file_size=1, content_hash="hash"
    ))
    user = SimpleNamespace(
        user_id="u1", display_name="张三", department_id="d1", department_name="研发部",
        group_id="g1", group_name="一组", role=ROLE_MINISTER,
    )
    early = datetime(2026, 7, 31, 8, 59, 59, tzinfo=ZoneInfo("Asia/Shanghai"))
    on_time = datetime(2026, 7, 31, 9, 0, tzinfo=ZoneInfo("Asia/Shanghai"))
    assert handler._new_report_date(early) == "2026-07-30"
    assert handler._new_report_date(on_time) == "2026-07-31"
    monkeypatch.setattr(handler, "_submission_now", lambda: early)
    monkeypatch.setattr("src.web.app.now_iso", lambda: "2026-07-31T08:59:59")
    report = handler._report_from_payload(user, {"date": "1999-01-01", "title_summary": "今日工作", "work_description": "完成日报功能"})
    assert report.report_date == "2026-07-30"
    assert report.created_at == "2026-07-31T08:59:59"


def test_business_user_cannot_choose_submission_time():
    handler = object.__new__(DailyReportHandler)
    handler.config = {"daily_report_submission": {"enabled": True, "timezone": "Asia/Shanghai", "rollover_time": "09:00"}}
    handler.db = SimpleNamespace(get_derived_manager=lambda user: None)
    handler.storage = SimpleNamespace(save_bytes=lambda *args: SimpleNamespace(
        original_filename="report.docx", stored_path="submitted_reports/report.docx", file_size=1, content_hash="hash"
    ))
    user = SimpleNamespace(
        user_id="u_test_employee", display_name="测试员工", department_id="", department_name="",
        group_id="", group_name="", role="employee", is_test_account=False,
    )
    with pytest.raises(ValueError, match="上海时间"):
        handler._report_from_payload(user, {"submitted_at": "2026-07-31T08:59", "manager": "组长", "title_summary": "今日工作", "work_description": "完成日报功能"})


def test_report_editing_window_respects_report_grouping():
    handler = object.__new__(DailyReportHandler)
    handler.config = {"daily_report_submission": {"enabled": True, "timezone": "Asia/Shanghai", "rollover_time": "09:00"}}
    from src.core.schemas import Report
    previous = Report("r1", "u1", "张三", "d1", "研发部", "2026-07-30", "组长", "标题", "内容")
    today = Report("r2", "u1", "张三", "d1", "研发部", "2026-07-31", "组长", "标题", "内容")
    tz = ZoneInfo("Asia/Shanghai")
    handler._ensure_report_editable(previous, datetime(2026, 7, 31, 8, 59, 59, tzinfo=tz))
    with pytest.raises(ValueError, match="超过"):
        handler._ensure_report_editable(previous, datetime(2026, 7, 31, 9, 0, tzinfo=tz))
    handler._ensure_report_editable(today, datetime(2026, 7, 31, 9, 0, tzinfo=tz))


def test_report_rollover_time_uses_the_saved_admin_setting():
    handler = object.__new__(DailyReportHandler)
    handler.config = {"daily_report_submission": {"enabled": True, "timezone": "Asia/Shanghai", "rollover_time": "09:00"}}
    handler.db = SimpleNamespace(get_settings=lambda defaults: defaults | {"daily_report_submission.rollover_time": "08:30"})
    now = datetime(2026, 7, 31, 8, 30, tzinfo=ZoneInfo("Asia/Shanghai"))
    assert flatten_settings(handler.config)["daily_report_submission.rollover_time"] == "09:00"
    assert handler._new_report_date(now) == "2026-07-31"
    with pytest.raises(ValueError, match="HH:MM"):
        handler._validate_settings_payload({"daily_report_submission.rollover_time": "9点"})


def test_automatic_daily_duplicate_settings_are_exposed_and_validated():
    settings = flatten_settings({"automatic_daily_duplicate": {"enabled": False, "run_at": "13:15"}})

    assert settings["automatic_daily_duplicate.enabled"] is False
    assert settings["automatic_daily_duplicate.run_at"] == "13:15"
    DailyReportHandler._validate_settings_payload({"automatic_daily_duplicate.enabled": True, "automatic_daily_duplicate.run_at": "13:15"})
    with pytest.raises(ValueError, match="true 或 false"):
        DailyReportHandler._validate_settings_payload({"automatic_daily_duplicate.enabled": "true"})
    with pytest.raises(ValueError, match="HH:MM"):
        DailyReportHandler._validate_settings_payload({"automatic_daily_duplicate.run_at": "13点15"})


def test_daily_duplicate_report_worker_count_is_exposed_and_validated():
    assert flatten_settings({"daily_duplicate": {}})["daily_duplicate.report_worker_count"] == 3
    DailyReportHandler._validate_settings_payload({"daily_duplicate.report_worker_count": 1})
    DailyReportHandler._validate_settings_payload({"daily_duplicate.report_worker_count": 8})
    for invalid in (0, 9, 2.5, "three"):
        with pytest.raises(ValueError, match="1 至 8"):
            DailyReportHandler._validate_settings_payload({"daily_duplicate.report_worker_count": invalid})


def test_local_worker_count_is_exposed_and_validated():
    assert flatten_settings({"daily_duplicate": {}})["daily_duplicate.local_worker_count"] == 3
    DailyReportHandler._validate_settings_payload({"daily_duplicate.local_worker_count": 1})
    DailyReportHandler._validate_settings_payload({"daily_duplicate.local_worker_count": 32})
    for invalid in (0, 33, 2.5, "many"):
        with pytest.raises(ValueError, match="1 至 32"):
            DailyReportHandler._validate_settings_payload({"daily_duplicate.local_worker_count": invalid})


def test_llm_retry_setting_matches_the_supported_implementation():
    DailyReportHandler._validate_settings_payload({"llm_judge.max_retries": 0})
    DailyReportHandler._validate_settings_payload({"llm_judge.max_retries": 1})
    with pytest.raises(ValueError, match="0 至 1"):
        DailyReportHandler._validate_settings_payload({"llm_judge.max_retries": 2})


def test_retired_worker_count_is_not_exposed_or_copied_into_new_tasks():
    db = SimpleNamespace(get_settings=lambda defaults: defaults | {"daily_duplicate.worker_count": 7})

    settings = effective_settings({"daily_duplicate": {}}, db)

    assert settings["daily_duplicate.report_worker_count"] == 3
    assert "daily_duplicate.worker_count" not in settings


def test_report_docx_marks_grouping_and_actual_submission_time():
    body = build_report_docx({"name": "张三", "department": "研发部", "report_date": "2026-07-30", "submitted_at": "2026-07-31T08:59:59", "manager": "组长", "title_summary": "标题", "work_description": "内容"})
    text = "\n".join(paragraph.text for paragraph in Document(BytesIO(body)).paragraphs)
    assert "查重归属日期：2026-07-30" in text
    assert "提交时间：2026-07-31T08:59:59" in text


def test_periodic_export_arguments_require_a_valid_date_range():
    assert _periodic_export_args({"report_type": ["weekly"], "start": ["2026-07-01"], "end": ["2026-07-31"]}) == ("weekly", "2026-07-01", "2026-07-31")
    import pytest
    with pytest.raises(ValueError, match="开始日期"):
        _periodic_export_args({"report_type": ["monthly"], "start": ["2026-08-01"], "end": ["2026-07-31"]})


def test_periodic_exports_include_complete_report_details():
    from src.core.schemas import Report

    report = Report("r1", "u1", "张三", "d1", "研发部", "2026-07-31", "组长", "完成接口", "完成接口设计、开发和联调。",
                    code_items=[{"file_name": "main.py", "language": "python", "work_type": "修改", "description": "新增接口"}],
                    document_items=[{"document_name": "设计说明", "document_type": "docx", "work_type": "新增", "description": "补充设计"}],
                    testcase_items=[{"用例编号": "TC-1", "描述": "接口测试"}], submitter_role="employee",
                    created_at="2026-07-31T09:00:00", updated_at="2026-07-31T10:00:00")
    artifacts = {"r1": [{"original_filename": "接口说明.pdf", "artifact_type": "document", "created_at": "2026-07-31T09:00:00"}]}

    word = _build_periodic_reports_docx("周报汇总", "本人", "2026-07-28", "2026-07-31", [report], artifacts)
    text = "\n".join(paragraph.text for paragraph in Document(BytesIO(word)).paragraphs)
    assert "完成接口设计、开发和联调。" in text
    assert "代码记录" in text
    book = load_workbook(BytesIO(_build_periodic_reports_xlsx("周报汇总", "本人", "2026-07-28", "2026-07-31", [report], artifacts)))
    assert book.sheetnames == ["日报明细", "代码记录", "文档记录", "测试用例", "附件"]
    assert book["日报明细"].max_row == 3
    assert book["附件"].max_row == 2


def test_periodic_word_export_handles_empty_reports():
    word = _build_periodic_reports_docx("月报汇总", "全公司", "2026-07-01", "2026-07-31", [], {})
    assert "暂无日报" in "\n".join(paragraph.text for paragraph in Document(BytesIO(word)).paragraphs)


def test_periodic_exports_mark_workday_without_report_as_none_and_skip_weekends():
    from src.core.schemas import Report

    submitter = {"user_id": "u1", "display_name": "张三", "role": "employee", "department_name": "研发部", "group_name": "一组"}
    db = SimpleNamespace(list_export_submitters=lambda user: [submitter])
    report = Report("r1", "u1", "张三", "d1", "研发部", "2026-07-31", "组长", "已提交", "工作内容", created_at="2026-07-31T09:00:00")
    missing = _periodic_missing_reports(db, SimpleNamespace(), "2026-07-30", "2026-08-02", [report])

    assert missing == [{"日期": "2026-07-30", "姓名": "张三", "职务": "employee", "部门": "研发部", "小组": "一组", "日报": "无"}]
    word = _build_periodic_reports_docx("周报汇总", "本人", "2026-07-30", "2026-08-02", [report], {}, missing)
    assert "无" in [cell.text for table in Document(BytesIO(word)).tables for row in table.rows for cell in row.cells]
    book = load_workbook(BytesIO(_build_periodic_reports_xlsx("周报汇总", "本人", "2026-07-30", "2026-08-02", [report], {}, missing)))
    rows = list(book["日报明细"].values)
    assert (None, "2026-07-30", "张三", "employee", "研发部", "一组", "无", "无", None, None, "无") in rows


def test_admin_task_control_audit_descriptors_and_allowed_actions():
    expected = {
        "pause": ("task.pause", "info"),
        "resume": ("task.resume", "info"),
        "restart": ("task.restart", "warning"),
        "force-terminate": ("task.force_terminate", "critical"),
    }
    for endpoint, (action, severity) in expected.items():
        descriptor = DailyReportHandler._audit_descriptor("POST", f"/api/admin/tasks/T_1/{endpoint}")
        assert descriptor["action"] == action
        assert descriptor["severity"] == severity
        assert descriptor["target_id"] == "T_1"

    handler = object.__new__(DailyReportHandler)
    handler.config = {"task_runner": {"stale_heartbeat_seconds": 90}}
    handler.db = SimpleNamespace(get_worker_lease=lambda: {"worker_id": "worker-1", "heartbeat_at": "2999-01-01T00:00:00"})
    response = handler._task_console_response({
        "tasks": [
            {"task_id": "q", "status": "queued", "effective_status": "queued"},
            {"task_id": "p", "status": "paused", "effective_status": "paused"},
            {"task_id": "f", "status": "failed", "effective_status": "failed"},
        ],
        "counts": {"queued": 1, "paused": 1, "failed": 1},
        "next_cursor": "",
    })

    assert response["worker"]["online"] is True
    assert response["counts"] == {"running": 0, "queued": 1, "paused": 1, "abnormal": 1}
    assert response["tasks"][0]["allowed_actions"] == ["pause", "force_terminate"]
    assert response["tasks"][1]["allowed_actions"] == ["resume", "restart", "force_terminate"]
    assert response["tasks"][2]["allowed_actions"] == ["restart"]
